# -*- coding: utf-8 -*-

"""Akvo RSR is covered by the GNU Affero General Public License.

See more details in the license.txt file located at the root folder of the
Akvo RSR module. For additional details on the GNU license please
see < http://www.gnu.org/licenses/agpl.html >.
"""

import base64
import binascii
import io
import re
import requests

from docx.document import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.image.exceptions import (
    UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError)
from docx.image.image import Image
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from markdown import markdown
from lxml.html import fromstring


def load_image(src):
    """ load image from base64 data or url
    """
    image_bytes = load_inline_image(src) \
        if src.startswith("data:") \
        else load_external_image(src)

    return make_image(image_bytes)


def load_external_image(src):
    image_data = None
    try:
        response = requests.get(src, stream=True)
        image_data = response.content
    except (requests.RequestException, IOError):
        pass

    return image_data


def load_inline_image(src):
    image_data = None
    header_data = src.split(";base64,", maxsplit=1)
    if len(header_data) == 2:
        data = header_data[1]
        try:
            image_data = base64.b64decode(data, validate=True)
        except (binascii.Error, ValueError):
            pass
    return image_data


def make_image(data):
    image_buffer = None
    if data:
        image_buffer = io.BytesIO(data)
        try:
            Image.from_blob(image_buffer.getbuffer())
        except (UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError):
            image_buffer = None

    if not image_buffer:
        broken_img = load_inline_image(IMG_BLANK)
        image_buffer = io.BytesIO(broken_img)

    return image_buffer


def set_repeat_table_header(row):
    """ set repeat table row on every page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

    return row


def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
    """ places a hyperlink within a paragraph object
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color is not None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def change_orientation(document):
    """ change document orientation from portrait to landscape and vice versa
    """
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE \
        if current_section.orientation == WD_ORIENT.PORTRAIT \
        else WD_ORIENT.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


def markdown_to_docx(container, text):
    """ insert markdown text to docx document
    """
    if not text:
        return

    builder = HTMLDocxBuilder(container)
    builder.feed(markdown(text))


class HTMLDocxBuilder(object):
    def __init__(self, container):
        self.root_container = container

    def feed(self, html):
        root = fromstring(html)
        self.traverse(self.root_container, root)

    def traverse(self, container, element):
        handler = get_tag_handler(element.tag)
        new_container = container
        if handler:
            new_container = handler.handle_text(container, element)
        for child in element:
            self.traverse(new_container, child)
        parent_handler = get_tag_handler(element.getparent().tag)
        if element.tail and parent_handler:
            parent_handler.handle_tail(container, element)


class ParagraphTagHandler(object):
    """ <p> creates a paragraph element inside a docx container element.
    """
    def handle_text(self, container, element):
        paragraph = get_new_paragraph(container)
        return self._append_paragraph(element.text, element, paragraph)

    def handle_tail(self, container, element):
        paragraph = get_current_paragraph(container)
        return self._append_paragraph(element.tail, element, paragraph)

    def _append_paragraph(self, text, element, container):
        text = trim_whitespaces(text)
        if not text:
            return container
        style = None
        if element.getparent().tag == 'blockquote':
            style = 'IntenseQuote'
        container.add_run(text=text, style=style)
        return container


class StrongTagHandler(object):
    """
    <strong> Creates a bold text run inside the paragraph container.
    Appends remainder of text as a additional run
    """
    def handle_text(self, container, element):
        return self._append_strong(element.text, element, container)

    def handle_tail(self, container, element):
        return self._append_strong(element.tail, element, container)

    def _append_strong(self, text, element, container):
        text = trim_whitespaces(text)
        run = container.add_run(text=text)
        run.bold = True
        if element.getparent().tag == 'em':
            run.italic = True
        return container


class EmphasisTagHandler(object):
    """
    <em> Creates an italic text run inside the paragraph container.
    Appends remainder of text as a additional run
    """
    def handle_text(self, container, element):
        return self._append_emphasis(element.text, element, container)

    def handle_tail(self, container, element):
        return self._append_emphasis(element.tail, element, container)

    def _append_emphasis(self, text, element, container):
        text = trim_whitespaces(text)
        run = container.add_run(text=text)
        run.italic = True
        if element.getparent().tag == 'strong':
            run.bold = True
        return container


class LineBreakTagHandler(object):
    """ <br> Creates a break item inside the given container.
    """
    def handle_text(self, container, element):
        element.tail = trim_whitespaces(element.tail)
        element.tail = element.tail.lstrip()
        run = container.add_run()
        run.add_break(break_type=WD_BREAK.LINE_CLEAR_RIGHT)
        return container

    def handle_tail(self, container, element):
        pass


class ListItemTagHandler(object):
    """
    <li> Create a list item element inside a docx container.
    Style it according to its parents list type.
    """
    def __init__(self):
        self.list_style = dict(
            ol='ListNumber',
            ul='ListBullet',
        )

    def handle_text(self, container, element):
        paragraph = get_new_paragraph(container)
        return self._append_list_item(element, element.text, paragraph)

    def handle_tail(self, container, element):
        paragraph = get_current_paragraph(container)
        return self._append_list_item(element, element.tail, paragraph)

    def _append_list_item(self, element, text, container):
        text = trim_whitespaces(text)
        text = '' if text == ' ' else text
        style = self.list_style.get(element.getparent().tag, 'ListBullet')
        container.style = style
        container.add_run(text)
        return container


class DivTagHandler(object):
    def handle_text(self, container, element):
        return self._append_run(element.text, container)

    def handle_tail(self, container, element):
        return self._append_run(element.tail, container)

    def _append_run(self, text, container):
        text = trim_whitespaces(text)
        text = '' if text == ' ' else text
        if text:
            container.add_run(text=text)
        return container


strong_handler = StrongTagHandler()
default_handler = DivTagHandler()
_tag_handler_map = dict(
    p=ParagraphTagHandler(),
    em=EmphasisTagHandler(),
    br=LineBreakTagHandler(),
    li=ListItemTagHandler(),
    strong=strong_handler,
    h1=strong_handler,
    h2=strong_handler,
    h3=strong_handler,
    h4=strong_handler,
    h5=strong_handler,
    h6=strong_handler,
)


def get_tag_handler(tag):
    return _tag_handler_map.get(tag, default_handler)


def get_current_paragraph(container):
    current_paragraph = container
    if isinstance(container, Paragraph):
        if isinstance(container._parent, _Cell):
            current_paragraph = container._parent.paragraphs[-1]

    if isinstance(container, Document):
        current_paragraph = container.add_paragraph()

    return current_paragraph


def get_new_paragraph(container):
    new_paragraph = container
    if isinstance(container, Paragraph):
        if isinstance(container._parent, _Cell):
            new_paragraph = container._parent.paragraphs[0]
            if len(container._parent.paragraphs) > 1:
                new_paragraph = container._parent.add_paragraph()
            else:
                if container._parent.paragraphs[0].text:
                    new_paragraph = container._parent.add_paragraph()
        else:
            if container.text:
                new_paragraph = container._parent.add_paragraph()

    if isinstance(container, Document):
        new_paragraph = container.add_paragraph()

    return new_paragraph


def trim_whitespaces(text):
    """ replaces multiple whitespaces and line breaks by a single whitespace
    """
    if text:
        text = ' '.join(text.split('\n'))
        text = re.sub(' +', ' ', text)

    return text if text else ''


IMG_BLANK = "data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCAJYAlgDASIAAhEBAxEB/8QAHQABAQADAQEBAQEAAAAAAAAAAAcBBggCBAUDCf/EAEcQAQABAgMCBg8GBAUEAwAAAAABAgMEBREGkwcSFyExcRMUMjU2QVFTVFVhdKGz0RYikrGy4UKBkcEjQ0VioxUzUvBygqL/xAAYAQEBAQEBAAAAAAAAAAAAAAAAAgMBBP/EACERAQEAAgEFAQEBAQAAAAAAAAABAhEDEhMxMlEhQUJh/9oADAMBAAIRAxEAPwD/AE1AWgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAJmImNebUmqKaZmZiIiNZmU72l4UasNjKsLlliK+JOk37mnPPs5+eOj4qxxuXgUOa6YmYiddPYceEbq4T8+mqrXE0Tz+b10Y5T8+9Io3S+3k5uLLx4OPCNcp+fekUbo5T8+9Io3R28jcWXjwceEa5T8+9Io3Ryn596RRujt5G4svHg48I1yn596RRujlPz70ijdHbyNxZePBx4RrlPz70ijdHKfn3pFG6O3kbiy8eDjwjXKfn3pFG6OU/PvSKN0dvI3Fl48HHhGuU/PvSKN0cp+fekUbo7eRuLLxtY5uefJrzs8/jiY60dscKWd266ZruWblGukxVa0+Lfdk9tcPtHTXarojD4qn+DXWK48sOXCx3bZQiddfYMwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB+Dt1ja8BspmF21Ol2aaaYnyRVVFP8AdE8Fgox2OweFt1zT2W5FuK641nWZ05/5ysnCP4I47rt/NoSTZ7v/AJV71R+ul6OP1qfN0oFHBBhooiasyvRVMazxaI01/q9ckOF9Z4j8EfVv/RFHtiI6uZ/DH46xldmq9i7tOHtR/HXPNPshn15fTUaNPBFhY/1LEfgj6kcEOFn/AFLER/8ASPq+nEcK+UUYiaaLOJvU083Hppp0nq+8/WyPbzKc9udiouThLuulNu/ERNXVpMqt5Ifj8HkhwvrPEfgj6nJDhfWeI/BH1b+I68vruo0DkhwvrPEfgj6nJDhfWeI/BH1b+HXl9NRoHJDhfWeI/BH1OSHC+s8R+CPq38OvL6ajQOSHC+s8R+CPqckOF9Z4j8EfVv4deX01GgckOF9Z4j8EfU5IcL6zxH4I+rfw68vpqJPtVwc28hym5jrWOuXotTTrRXTEa61RTHxlrOzOMqwW0WAvW+aubsUT/OdFZ4Ro12Qx3Xb+bQj+SxpneXe8UfnLbC3LHdTPLoarSK6op7mJ5tWGZ7urrYearAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAazwj+COO67fzaEk2e7/AOVe9UfrpVvhH8Ecd12/m0JJs93/AMq96o/XS34/VM8ugqdZ0iPHFMaz4kU2/wBobue5zfs9kq7Us18SijxTp41rpmPux0a088z1Oes3szhszxNqvWmum5Vrr1zJxyW2u5eHx6c+uuv8tNHqiuq3et3aK5oqomZiYjnefj7TX+vkei/9ZrbsFn9WfZDTVd/7+Hqi1XPjq8cVfH4Nj8cx5J0aFwR2KreWY+9rrRcvRbp9ukROsf1+DftIiIiHiymsq1jACQZ016OeZ6I8rzVXTREzVMU00xrVVPRTHllLdvdvK8dXOAy25VRhon/FuxzTcmJ5tJjxR9VY49VNxUuNEzzTr7Xpo2w+3dObW7eDzCri4+OaLukRTd8nj6f2bxM6TpMTFUdMT4nLNXRvYMsODW+EXwPx3Xa+ZQj+Td+8u94o/OVg4RfA/Hddr5lCP5N37y73ij85b8fqmeXQs93V1sMz3dXWwwqgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGs8I/gjjuu382hJNnu/8AlXvVH66Vb4R/BHHddv5tCSbPd/8AKveqP10t+P1TPLoGYiqLcT0RpM/05mg7fbB380v1ZlllPHxVXNdsx448sf8Avib94o6o/J6408Xi8aYjyQxluN3FeXOV/B4jCXptX7VdFyJ7mqNJh+pkeyGY7QX6abViabOulV2rmj2x/RdqqaZjmoo18vEpmfjDOmkRE18SmebXSIjX+UNry7n5E9MfFlGUWcky+xg7FPFt2qdPbM+V90vNMTEc9M0z5J6WWHlQxVXTREzVMU00xrVVPRTHlkqrpoiZqmKaaY1qqnopjyylm3+3dWYcbLsurmnD0z9+7H+ZPjiPZ0fFWONyrlpt/t3VmHGy7Lq5pw9M/fux/mT44j2dHxaHERHRzewiIjo5vYPZjjMYzeqa6rVdFy3VNFymdaao/h9sKtsRt3RmtFGCx1U04yn7tFydNLkc2nPr3XT8Eneqa6rVdFy3VNFymdaao/h9sOZYzKEuvDo6Z0nSYmKo6YnxMtG2I27ozWijBY6qacZT92i5OmlyObTn17rp+DeJnSdJiYqjpifE8dll1Wsu2t8Ivgfjuu18yhH8m795d7xR+crBwi+B+O67XzKEfybv3l3vFH5y24/VM8uhZ7urrYZnu6uthhVAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAANZ4R/BHHddv5tCSbPd/8AKveqP10q3wj+COO67fzaEk2e7/5V71R+ulvx+qZ5dA+KOqPyDxR1R+R01U0xGtVXNTHl8rCqOmqmmI1qq5qY8vlaBt3wgTgZnL8srpqvxOt2/HPFPP3Me36sbe7e9qU3cty27/jVxxb1+jn4sf8AjE+VMJnWZno9muvxejDD+1y1adjNtLe01iLN7SjMLdP3qfORH8Uf++Jsk3KYp481RTbiONNc80RHllzthsRcwd+i/Zrqt37c8a3cpnSaJ8vtbRtDwi5hnuVWsJERYq0ib1dHN2Sfp0GXH+/jm/x9+3+3dWYcbLsurmnD0z9+7H+ZPjiPZ0fFocREdHN7CIiOjm9g2xxmMQAKAAHqmuq1XRct1TRcpnWmqP4fbCt7A7Y1Z9ROCxNNdWLs0TX2XSOLVRGnTOuuusykTeOCWNdocTHlwtUf/qmP7ss5LNuz8brwjeCGPjyTbj/lpR7Ju/eXe8UfnKwcIs67IY+fLVbn/lpR/Ju/eXe8UfnKOP1V/p0LPd1dbDM93V1sPPVAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAANZ4R/BHHddv5tCSbPd/8AKveqP10q3wj+COO67fzaEk2e7/5V71R+ulvx+qZ5dAx/DEzpzRzz1J/t7t52pRdy3Lbkxdrji3r9Pip8lPtb/XEdin/4THwc7Y7nx2JmembtWv8AXQ48d12/j+UzrMz0ezXX4sA9LMAAAAAAAAbxwR+EV/3af10NHbxwR+EV/wB2n9dCM/Wuzy3PhE8D8d12vmUJBk3fvLveKPzlX+ETwPx3Xa+ZQkGTd+8u94o/OWXH6q/06Fnu6uthme7q62HnqgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGs8I/gjjuu382hJNnu/+Ve9UfrpV3hFt1VbH46aY42nE5o6eaumfyhH8lvW8NnOXXblcU0W79Ncz7IqiXo4/Wp8V0HXz2vbpHN5eZG8VwdZ9cxV+ujBRNNVyqqJm5TGsTOvjlU6NpsnuURP/UsL0R92q7FOnN7WftHlPrPCz136Wcyyx8R3cqT8m+0HoUb6j6nJvtB6FG+o+qsfaPKfWWE39B9o8p9ZYTf0K7mfxzUSfk32g9CjfUfU5N9oPQo31H1Vj7R5T6ywm/oPtHlPrLCb+g7mfw1En5N9oPQo31H1OTfaD0KN9R9VY+0eU+ssJv6D7R5T6ywm/oO5n8NRJ+TfaD0KN9R9Tk32g9CjfUfVWPtHlPrLCb+g+0eU+ssJv6DuZ/DUSfk32g9CjfUfU5N9oPQo31H1Vj7R5T6ywm/oPtHlPrLCb+g7mfw1En5N9oPQo31H1bRwd7KZpkOdXb2Nw3YrdVmaImK6Z5+NTPin/bLcPtHlPrLCb+g+0eU+scFM/wC69S5c8rNWH4/M4RObZDHx5JtR/wAlCP5N37y73ij85U/b/aDLb+zGJsWsdYvXbtVHFptVROmldNXi9kJjkkTXnuXU0xMz2amrSPHpK+OXpP66Enu6uthmrmrq62HnqgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHjEYW3jsNew16ia7d2iaZj8kP2h2Qx+z9+Yrw9VdrpprtxrTHPPxXOYidNY6J111KqYrn72ldOvc1RrC8crieXN8W9ZnW3VE68/+Hr/AGZ7HHm6t1+zoztezrOlmzpM+bg7XteZs7qG05XNRzn2OPN1br9jscebq3X7OjO17XmbO6g7XteZs7qDvGo5z7HHm6t1+x2OPN1br9nRna9rzNndQdr2vM2d1B3jUc59jjzdW6/Y7HHm6t1+zozte15mzuoO17XmbO6g7xqOc+xx5urdfsdjjzdW6/Z0Z2va8zZ3UHa9rzNndQd41HOfY483Vuv2Oxx5urdfs6M7XteZs7qDte15mzuoO8ajnPscebq3X7HY483Vuv2dGdr2vM2d1B2va8zZ3UHeNRzraw1d67TTas113Ne5i3p/ZRtgNiMRh8bGZZhb7H2OdbNuvutdI59OjT+filRKbFqJ17FbjqtxD+lU66adEIvJf4a0x450AYugAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAP/9k="
