# -*- coding: utf-8 -*-

"""Akvo RSR is covered by the GNU Affero General Public License.

See more details in the license.txt file located at the root folder of the
Akvo RSR module. For additional details on the GNU license please
see < http://www.gnu.org/licenses/agpl.html >.
"""

import os

from akvo.rsr.models import Project, IndicatorPeriod, ProjectUpdate
from akvo.rsr.decorators import with_download_indicator
from datetime import datetime
from dateutil.relativedelta import relativedelta
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import utils
from .docx_utils import load_image, add_hyperlink, set_repeat_table_header, change_orientation, markdown_to_docx


def build_view_object(project, start_date=None, end_date=None):
    periods = IndicatorPeriod.objects\
        .select_related('indicator', 'indicator__result', 'indicator__result__project')\
        .filter(indicator__result__project=project)
    if start_date and end_date:
        periods = periods.filter(
            Q(period_start__isnull=True) | Q(period_start__gte=start_date),
            Q(period_end__isnull=True) | Q(period_end__lte=end_date)
        )

    if not periods.count():
        return utils.ProjectProxy(project)

    return utils.make_project_proxies(periods.order_by('-period_start'))[0]


def get_project_updates(project, start_date=None, end_date=None):
    updates = ProjectUpdate.objects.filter(project=project)
    if start_date and end_date:
        updates = updates.filter(event_date__gte=start_date, event_date__lte=end_date)

    return [utils.ProjectUpdateProxy(u) for u in updates.order_by('-created_at')]


def is_empty_value(value):
    if not value:
        return True
    if value == '0':
        return True
    if value == 'N.A.':
        return True


def build_log_frame(project_view):
    data = []
    use_baseline = False
    has_disaggregations = False
    previous_result = ''
    previous_indicator = ''
    for result in project_view.results:
        for indicator in result.indicators:
            if indicator.is_qualitative:
                break
            if not is_empty_value(indicator.baseline_value):
                use_baseline = True
            for period in indicator.periods:
                current_result = ''
                if previous_result != result.title:
                    previous_result = result.title
                    current_result = result.title
                current_indicator = ''
                if previous_indicator != indicator.title:
                    previous_indicator = indicator.title
                    current_indicator = indicator.title
                disaggregations = []
                for d in period.disaggregations.all():
                    disaggregations.append({
                        'label': str(d.dimension_value),
                        'value': d.value
                    })
                if len(disaggregations):
                    has_disaggregations = True
                    print(disaggregations)

                data.append({
                    'result': current_result,
                    'type': result.iati_type_name,
                    'indicator': current_indicator,
                    'baseline_year': indicator.baseline_year if current_indicator != '' else '',
                    'baseline_value': indicator.baseline_value if current_indicator != '' else '',
                    'indicator_target': '{:,}'.format(indicator.target_value) if current_indicator != '' else '',
                    'period_start': period.period_start,
                    'period_end': period.period_end,
                    'target_value': int(period.target_value),
                    'actual_value': int(period.actual_value),
                    'comments': period.actual_comment,
                    'disaggregations': disaggregations,
                })

    return {
        'data': data,
        'use_baseline': use_baseline,
        'use_indicator_target': project_view.use_indicator_target,
        'has_disaggregations': has_disaggregations,
    }


def prepare_result_title(iati_type, title):
    if not title:
        return ('', '')

    if not iati_type:
        return ('', title)

    head, separator, tail = title.strip().partition(':')
    if tail and head.split()[0] == iati_type:
        return ('{}: '.format(head.strip()), tail.strip())

    return ('{}: '.format(iati_type), title.strip())


@login_required
@with_download_indicator
def render_report(request, project_id):
    project = get_object_or_404(Project, pk=project_id)
    start_date = utils.parse_date(request.GET.get('start_date', '').strip(), datetime(1900, 1, 1))
    end_date = utils.parse_date(request.GET.get('end_date', '').strip(), datetime.today() + relativedelta(years=10))

    project_view = build_view_object(project, start_date, end_date)
    project_updates = get_project_updates(project, start_date, end_date)
    log_frame = build_log_frame(project_view)
    today = datetime.today()

    if request.GET.get('show-html', ''):
        html = render_to_string('reports/project-kickstart.html', context={
            'project': project_view,
            'project_updates': project_updates,
            'log_frame': log_frame
        })

        return HttpResponse(html)

    doc = Document(os.path.join(os.path.dirname(__file__), 'kickstart.tpl.docx'))
    doc.sections[0].page_width = Mm(210)
    doc.sections[0].page_height = Mm(297)
    doc.sections[0].footer.paragraphs[-1].text = 'Akvo RSR report {}'.format(today.strftime('%Y-%m-%d'))

    doc.add_heading(project_view.title, 0)
    doc.add_paragraph(today.strftime('%Y-%m-%d'), 'Subtitle')
    doc.add_page_break()

    if project_view.project_plan.strip():
        doc.add_heading('Project plan', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.project_plan.strip())
    if project_view.goals_overview.strip():
        doc.add_heading('Goals overview', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.goals_overview.strip())
    if project_view.target_group.strip():
        doc.add_heading('Target group', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.target_group.strip())
    if project_view.project_plan_summary.strip():
        doc.add_heading('Summary of project plan', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.project_plan_summary.strip())
    if project_view.background.strip():
        doc.add_heading('Background', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.background.strip())
    if project_view.current_status.strip():
        doc.add_heading('Situation at start of project', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.current_status.strip())
    if project.sustainability.strip():
        doc.add_heading('Sustainability', 1)
        markdown_to_docx(doc.add_paragraph(), project_view.sustainability.strip())

    doc.add_heading('Project partners', 1)
    partners_table = doc.add_table(rows=1, cols=2)
    partners_table.style = 'Table Common'
    partners_table.rows[0].cells[0].paragraphs[-1].add_run('Organisation name').bold = True
    partners_table.rows[0].cells[1].paragraphs[-1].add_run('Roles').bold = True
    for partner in project_view.partnerships.all():
        row = partners_table.add_row()
        cell_p = row.cells[0].paragraphs[-1]
        if partner.organisation.url:
            add_hyperlink(cell_p, partner.organisation.url, partner.organisation.name)
        else:
            cell_p.text = partner.organisation.name
        row.cells[1].text = partner.iati_organisation_role_label_unicode()

    doc.add_paragraph('')

    doc.add_heading('Project budget', 1)
    budget_table = doc.add_table(rows=1, cols=5)
    budget_table.style = 'Table Common'
    budget_table.rows[0].cells[0].paragraphs[-1].add_run('Label').bold = True
    budget_table.rows[0].cells[1].paragraphs[-1].add_run('Period start').bold = True
    budget_table.rows[0].cells[2].paragraphs[-1].add_run('Period end').bold = True
    budget_table.rows[0].cells[3].paragraphs[-1].add_run('Amount').bold = True
    budget_table.rows[0].cells[4].paragraphs[-1].add_run('Currency').bold = True
    for budget in project.budget_items.all():
        row = budget_table.add_row()
        row.cells[0].text = budget.get_label()
        row.cells[1].text = budget.period_start.strftime('%Y-%m-%d') if budget.period_start else ''
        row.cells[2].text = budget.period_end.strftime('%Y-%m-%d') if budget.period_end else ''
        row.cells[3].text = '{:,}'.format(budget.amount)
        row.cells[4].text = budget.get_currency()
    for num, (currency, amount) in enumerate(project.budget_currency_totals().items()):
        row = budget_table.add_row()
        if num == 0:
            row.cells[0].paragraphs[-1].add_run('Total').bold = True
        row.cells[3].paragraphs[-1].add_run('{:,}'.format(amount)).bold = True
        row.cells[4].paragraphs[-1].add_run(currency).bold = True

    doc.add_paragraph('')

    doc.add_heading('Summary of results', 1)
    markdown_to_docx(
        doc.add_paragraph(),
        'Detail of overview of results are presented in **Appendix: Results log frame**.'
        'Here, we have summarised the results in terms of progress percentage.'
    )

    prog_form = load_image(IMG_PROGRESS_FORMULA)
    doc.add_picture(prog_form, width=Mm(117.5))

    legend_table = doc.add_table(rows=4, cols=3)
    legend_table.style = 'Table Common'

    legend_table.rows[0].cells[0].paragraphs[-1].add_run('Legend').bold = True
    legend_table.rows[0].cells[0].width = Mm(20)
    legend_table.rows[0].cells[1].width = Mm(30)
    legend_table.rows[0].cells[2].width = Mm(110)

    legend_table.rows[1].cells[0].paragraphs[-1].add_run().add_picture(load_image(IMG_GRADE_HIGH))
    legend_table.rows[1].cells[1].text = '85% - 100%'
    legend_table.rows[1].cells[2].text = 'Result nearly reached or reached'
    legend_table.rows[1].cells[0].width = Mm(20)
    legend_table.rows[1].cells[1].width = Mm(30)
    legend_table.rows[1].cells[2].width = Mm(110)

    legend_table.rows[2].cells[0].paragraphs[-1].add_run().add_picture(load_image(IMG_GRADE_MEDIUM))
    legend_table.rows[2].cells[1].text = '50% - 84%'
    legend_table.rows[2].cells[2].text = 'Result partly reached'
    legend_table.rows[2].cells[0].width = Mm(20)
    legend_table.rows[2].cells[1].width = Mm(30)
    legend_table.rows[2].cells[2].width = Mm(110)

    legend_table.rows[3].cells[0].paragraphs[-1].add_run().add_picture(load_image(IMG_GRADE_LOW))
    legend_table.rows[3].cells[1].text = '0% - 49%'
    legend_table.rows[3].cells[2].text = 'Result under reached'
    legend_table.rows[3].cells[0].width = Mm(20)
    legend_table.rows[3].cells[1].width = Mm(30)
    legend_table.rows[3].cells[2].width = Mm(110)

    doc.add_paragraph('')

    quantitative_table = doc.add_table(rows=1, cols=3)
    quantitative_table.style = 'Table Common'

    quantitative_table.cell(0, 0).merge(quantitative_table.cell(0, 2))
    title_p = set_repeat_table_header(quantitative_table.rows[0]).cells[0].paragraphs[-1]
    title_p.text = 'Results and indicators (quantitative)'
    title_p.style = 'Heading 2'
    quantitative_table.rows[0].cells[0].width = Mm(160)

    for result in project_view.results:
        if not result.has_quantitative_indicators:
            continue
        row = quantitative_table.add_row()
        row.cells[0].merge(row.cells[2])
        result_title_p = row.cells[0].paragraphs[-1]
        result_title_p.text = result.title
        result_title_p.style = 'Heading 3'
        row.cells[0].width = Mm(160)

        for indicator in result.indicators:
            if not indicator.is_quantitative:
                continue
            row = quantitative_table.add_row()
            row.cells[0].text = indicator.title
            progress_p = row.cells[1].paragraphs[-1]
            progress_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            progress_p.add_run(indicator.progress_str).bold = True
            grade_image = IMG_GRADE_HIGH if indicator.grade == 'high' \
                else IMG_GRADE_MEDIUM if indicator.grade == 'medium' \
                else IMG_GRADE_LOW
            row.cells[2].paragraphs[-1].add_run().add_picture(load_image(grade_image))
            row.cells[0].width = Mm(120)
            row.cells[1].width = Mm(20)
            row.cells[2].width = Mm(20)

    doc.add_paragraph('')

    doc.add_heading('Results and indicators (qualitative)', 2)
    for result in project_view.results:
        if not result.has_qualitative_indicators:
            continue
        for indicator in result.indicators:
            if not indicator.is_qualitative:
                continue
            doc.add_heading(indicator.title, 3)
            for period in indicator.periods:
                if not period.has_qualitative_data:
                    continue
                doc.add_paragraph('{} to {}'.format(
                    period.period_start.strftime('%Y-%m-%d'), period.period_end.strftime('%Y-%m-%d')))
                for update in period.approved_updates:
                    markdown_to_docx(doc.add_paragraph(), update.narrative)
                    if update.photo:
                        doc.add_paragraph().add_run().add_picture(load_image(update.photo_url), width=Mm(100))
                    if update.file:
                        add_hyperlink(doc.add_paragraph(), update.file_url, update.file_url)
                    doc.add_paragraph('\n\n')

    doc.add_heading('Project updates', 1)
    for project_update in project_updates:
        doc.add_heading(project_update.title, 2)
        doc.add_paragraph(project_update.created_at.strftime('%Y-%m-%d'))
        doc.add_paragraph().add_run().add_picture(load_image(project_update.photo_url), width=Mm(100))
        markdown_to_docx(doc.add_paragraph(), project_update.text)
        doc.add_paragraph('\n\n')

    new_section = change_orientation(doc)
    new_section.left_margin = Mm(10)
    new_section.top_margin = Mm(15)
    new_section.right_margin = Mm(10)
    new_section.bottom_margin = Mm(20)

    doc.add_heading('Appendix: Results log frame', 1)

    cols = 5
    if log_frame['use_baseline']:
        cols += 1
    if log_frame['use_indicator_target']:
        cols += 1
    if log_frame['has_disaggregations']:
        cols += 1

    log_table = doc.add_table(rows=1, cols=cols)
    log_table.style = 'Table Common'
    log_table.rows[0].cells[0].paragraphs[-1].style = 'Normal Smaller'
    log_table.rows[0].cells[0].paragraphs[-1].add_run('Result title').bold = True
    log_table.rows[0].cells[1].paragraphs[-1].style = 'Normal Smaller'
    log_table.rows[0].cells[1].paragraphs[-1].add_run('Indicator title').bold = True
    cell = 2
    if log_frame['use_baseline']:
        log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
        log_table.rows[0].cells[cell].paragraphs[-1].add_run('Baseline').bold = True
        cell += 1
    if log_frame['use_indicator_target']:
        log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
        log_table.rows[0].cells[cell].paragraphs[-1].add_run('Target').bold = True
        cell += 1
    log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
    log_table.rows[0].cells[cell].paragraphs[-1].add_run('Periods').bold = True
    cell += 1
    log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
    log_table.rows[0].cells[cell].paragraphs[-1].add_run('Values').bold = True
    cell += 1
    log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
    log_table.rows[0].cells[cell].paragraphs[-1].add_run('Comments').bold = True
    cell += 1
    if log_frame['has_disaggregations']:
        log_table.rows[0].cells[cell].paragraphs[-1].style = 'Normal Smaller'
        log_table.rows[0].cells[cell].paragraphs[-1].add_run('Disaggregations').bold = True

    set_repeat_table_header(log_table.rows[0])

    for log in log_frame['data']:
        row = log_table.add_row()
        result_head, result_body = prepare_result_title(log['type'], log['result'])
        if result_body:
            row.cells[0].paragraphs[-1].add_run(result_head).bold = True
            row.cells[0].paragraphs[-1].add_run(result_body)
        row.cells[0].paragraphs[-1].style = 'Normal Smaller'
        row.cells[1].text = log['indicator']
        row.cells[1].paragraphs[-1].style = 'Normal Smaller'
        cell = 2
        if log_frame['use_baseline']:
            row.cells[cell].paragraphs[-1].text = 'Year:\n{}'.format(log['baseline_year'])
            row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
            row.cells[cell].add_paragraph('Value:\n{}'.format(log['baseline_value']), 'Normal Smaller')
            cell += 1
        if log_frame['use_indicator_target']:
            row.cells[cell].paragraphs[-1].text = log['indicator_target']
            cell += 1
        row.cells[cell].paragraphs[-1].text = 'Start:\n{}'.format(
            log['period_start'].strftime('%Y-%m-%d') if log['period_start'] else '')
        row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
        row.cells[cell].add_paragraph(
            'End:\n{}'.format(log['period_end'].strftime('%Y-%m-%d') if log['period_end'] else ''),
            'Normal Smaller')
        cell += 1
        if not log_frame['use_indicator_target']:
            row.cells[cell].paragraphs[-1].text = 'Target:\n{:,}'.format(log['target_value'])
            row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
            row.cells[cell].add_paragraph('Actual:\n{:,}'.format(log['actual_value']), 'Normal Smaller')
        else:
            row.cells[cell].paragraphs[-1].text = '{:,}'.format(log['actual_value'])
            row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
        cell += 1
        row.cells[cell].text = log['comments']
        row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
        cell += 1
        if log_frame['has_disaggregations']:
            for k, d in enumerate(log['disaggregations']):
                if k == 0:
                    row.cells[cell].paragraphs[-1].text = '{}:\n{}'.format(d['label'], d['value'])
                    row.cells[cell].paragraphs[-1].style = 'Normal Smaller'
                else:
                    row.cells[cell].add_paragraph('{}:\n{}'.format(d['label'], d['value']), 'Normal Smaller')

    filename = '{}-{}-kickstart-report.docx'.format(today.strftime('%Y%b%d'), project.id)

    return utils.make_docx_response(doc, filename)


IMG_GRADE_HIGH = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAADEAAAAPCAYAAABN7CfBAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsQAAA7EAZUrDhsAAAJ3SURBVEhLtZa/bhNBEMZn92zHF+MIaBBJg4SIhESFkKgQHQ2IGkRNQYvEA/ACqSkpqVFeIBJlSIFASEGhIkEpCOKS4H93u8y3mcvd2Wd71xK/xjOru/V9O7Mzo4wxlkrY0YjMYEBmOJQVIqU1qWaTdBw7eybWkBkNeJ8em6ks8h4q4j3apFsxNpTVeo753c39bdr8sU3fjg9klehqfIkerd2hJ9fuUbfJ+wiFCGspOz2tfHwdUadDemlJvCo2S8n0/vBWmazUwAKi+CKpqCELVXaTA3q185Z+9n7LyiQXGm16c/cFra+sOv9cRJokZNPi5GYRcUQQlQomo/Qv/zFHYi5ThODDn33YoJO0LyvTKQtxcTU9Dr2nAJDh+ax62lk/8RMAkHIcsXFef3rnJQDgOUQMnIngOxAKhOfYdMiiRuL5gZSzo+KDPx7t0c7Rd/H8QORwdzQusjWeJ1iifHdsGn4IwJTe2zr8IlYYW4efWURAGo2Tp1S5CgVRSr/dZF+sMFDJ5tTLOSwQwf+BVo36UufD+btz6r4P5bofCte6SMwwXNNT6syOWu43FKWLA7x9+bpYYayvrHEk0I0XiEa5T+gmN78FoqFby2KR68So/aE85e7t/hldeO44UQKiK12bBehWRxw/NNJHF1mAdHp+44F4fmD8wCjivlxxSnnNRQwENLpd8QowE2mejXzAc7o9uQdO9SFHxAc89/LmY2dXBkCUTMxP08ouTj9a5hSQu1AHeobpn9TPTxIxNwTOALV/4+v72vkJKYeIQXDOxBQLIMYJyUso7g1PsSEpN97FMSepRv3gOA3XxX/tiYdLvEr3r9wSL4foH3/s/NHkH+75AAAAAElFTkSuQmCC"

IMG_GRADE_MEDIUM = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAADEAAAAPCAYAAABN7CfBAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsQAAA7EAZUrDhsAAAJ3SURBVEhLxZbPbtNAEMbH6/hP0iZNxSUCVeIYblyQOCLBA5QnoA/AA/AEwAPwApQ79MiRYwWX3ODYC1IkhNqmNLFje818m3FjN3FimwM/ydrd0a613+7Mzlha65RukSYJpfO5jBilyHIcsritgg4mND37SuH4h1iInP492h0+JeX3xLIZnaYUpwnFOhELb8NS5KoWt5ZYFhREpFFEyfU1pVqLpYhyXbJ3dnhV8Sd5zk/f0+8v71jIlViWKL9L+49f0J0nL8Wynlkyp2kckqaV8yU+TmrbLnVanlhyInQYGgHbwG3Y3S5Zti2WJeOTVzQZncionPb9R3Rw9EFGRa6iGQU6klE5jmVT3+UDZYx/wH2qCAC4JTOXrzsPTr+KADA7+0a/Pr+W0RKcfhUBIGJX+xMHpm9EVBWQkcYxu8viBwAxADeqw/npMUUXP2W0iAG4UR0wX6eaFE4Wm6oL3C/jcvRpbQxs4yInPOQbWBcD25iyEIVgboIJfnGpGb9ETQjG36XH7qHrHyTA66Wo5CWqQnaDCbvTvwJ3aopCDmjMmhfqf8D5o5kIrMvW+oMHpq2LnUt8SGJNQOJTdTJxHqzL6D18Lr167A6fSY9F2M1EuMpZPLGq3TaGOuTXeHwTSGB18AbDgvgWJy8ksDpgjW9nIjzPfFVB6XH79gaHb01ZUQXMGxy+kdGSntPhDZWXNHkwr+v40hewsW1CsPGyeSjwDo6OzQlvwunflXmrcQT/3udSAie8CZsF7Lmdm3krVSzyhuYKFm1WCFqtlvmU71eKnwknv8vRR1NeZMDdEAN77EJVKtkgifibm/IiA+7msft4iIObIpToLzbK9HBv1L7MAAAAAElFTkSuQmCC"

IMG_GRADE_LOW = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAADEAAAAPCAYAAABN7CfBAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsQAAA7EAZUrDhsAAAJiSURBVEhLxZa9UttAEMdXp09Lwc5MGKeFtFDzElDD21BQ8QD0vAQ9PT1tIDOZFExmMji29X2X3fWdJSuyIqnJTyPr9izZ+u/u3a4lpVTQIHt7g8XTE5SrFdvefA7TszOwo4jtPqgiA1Xm2gKwbBcsx9NWPzJZQI6nwbFs8PF3muyISF5e4Mf9Payen/XMLofn5zC/uuoUo4oUZLIEpUo9U8MSILwIz4meaCdF8csigVJJPVNh4RE5PoR4GrYi6MW/3d5CuV7zF/sIjo7gy81NqxCZLEDmibb2I9wARDDV1i6LPIakzLS1n8D2YOpunCHog9KnjwAieX2Fr9fX2qqQWdxLAEH3yXSTqnXo5fsIIOi+pf4/FvH97q6XAAMJ+fX4qC0Ewy6zv1+qC75fVimnlILfPZ1gWJcpp5ygKOxbA138fHjQI/JsykKGIrPKcanMQeExlBg3EDFGAEHRMKieKdBE1XaetKzGQ8hxA+FIjGXrgBFRaCJHRMHAa2Isdhjq0f9FfDg91cPhBMfHfLWEw9fBYN0wuFjIxiCwbojo5GSUR6mCG6xa4RmCqD3n2eMc4eNz7IpPFxc8MYTD2jPUTlBbMQQLPW9h0TN4GE13YERtjCQVPRbxGVsJqsR9mV9eAkWwjvAPdtKjE2o/JjNtVBw4AbcVfZm5mwzath3U7FHRo8avCxJAottQuE3K+L29bzKgAHvyESPX7vUCC+B7vm7tmwwkdOaFHD22m10siaBqXBdDa4bWADV/1NF2wtU75kawXgdM+nDz94+IUfWOuQXJoag5ZJM+LoS2j79nIgbwBx9dD8OwFXfOAAAAAElFTkSuQmCC"

IMG_PROGRESS_FORMULA = "data:image/jpg;base64,/9j/4AAQSkZJRgABAQAASABIAAD/4QBMRXhpZgAATU0AKgAAAAgAAYdpAAQAAAABAAAAGgAAAAAAA6ABAAMAAAABAAEAAKACAAQAAAABAAAC2aADAAQAAAABAAAASAAAAAD/7QA4UGhvdG9zaG9wIDMuMAA4QklNBAQAAAAAAAA4QklNBCUAAAAAABDUHYzZjwCyBOmACZjs+EJ+/8AAEQgASALZAwEiAAIRAQMRAf/EAB8AAAEFAQEBAQEBAAAAAAAAAAABAgMEBQYHCAkKC//EALUQAAIBAwMCBAMFBQQEAAABfQECAwAEEQUSITFBBhNRYQcicRQygZGhCCNCscEVUtHwJDNicoIJChYXGBkaJSYnKCkqNDU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6g4SFhoeIiYqSk5SVlpeYmZqio6Slpqeoqaqys7S1tre4ubrCw8TFxsfIycrS09TV1tfY2drh4uPk5ebn6Onq8fLz9PX29/j5+v/EAB8BAAMBAQEBAQEBAQEAAAAAAAABAgMEBQYHCAkKC//EALURAAIBAgQEAwQHBQQEAAECdwABAgMRBAUhMQYSQVEHYXETIjKBCBRCkaGxwQkjM1LwFWJy0QoWJDThJfEXGBkaJicoKSo1Njc4OTpDREVGR0hJSlNUVVZXWFlaY2RlZmdoaWpzdHV2d3h5eoKDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uLj5OXm5+jp6vLz9PX29/j5+v/bAEMAAgICAgICAwICAwUDAwMFBgUFBQUGCAYGBgYGCAoICAgICAgKCgoKCgoKCgwMDAwMDA4ODg4ODw8PDw8PDw8PD//bAEMBAgICBAQEBwQEBxALCQsQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEP/dAAQALv/aAAwDAQACEQMRAD8A/fyiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKa54oAN3J9qMmvmL4i/GLxXcfFG1+AXwYtbO68XvZjVNW1HUVkk07QtOd9kck0UTRvcXM7ZENuskeQC7yIgGfd7/V7bwj4Vutf8VXyC20aye6v7vZ5abLeMvNLtBbauFZsZOOmT1oA6cHNLXwP+xn+3j4U/bA1LxXo2meG5/Ct94YS2lWG4uVuWuYLhpELrtRNvlsihgc8uMGvvUNnnFAD6KjL4FeL/CX9oH4VfG7UPFOm/DXV21Sfwbe/wBn6mDbzwCG53Ou1WlRBIMxsNyFhx15oA9sopu7vSb6AH0Uzcc/WvDf2kPiX4q+EHwX8UfFDwfplnrN54XtX1CW1vJpII3toAWm2tGjneFBKggAkYJHWgD3XIor5Z/Y5+P+r/tM/AnSfi7rWkwaLd6lc3sLWttI0kaC2naIEM+CSwGTX1Hv5xxmgB9FR7/XFG8DqelAElGR0rkfHep+ItG8Hazq3hG2tr3WLK1lntoLyR4reWSNd2x5I1dlDAEZCnB7V8hfsHftb67+154B8R+Mtd0C18PyaJqf2BIbWV5VdPJjl3sXAwfnxgccUAfddFNDE496bvzz6elAElFM3HJA7UoYYz29aAHUUZFGRQAUZprdOma/OP4qft6a58N/2j7b9mW0+Flxr3iPVTEdMki1e3giu4Z1ZkYmWPEZ+RgVZuCMZIIJAP0eoyK/O+w/b9svDnx10L9nr47/AA51T4ceJfExt006WS7tdSs5Wu5GhgzNbNwskimMMAcPw4UAsP0MBwfegCSimBs9qN44x1NAD6KZuP4U0tkdj1H40AS0V+c2k/tkfEE/t1y/sheIPC2m2+n7J7iLVLa5mkla3Fi19AWjZFUOVwsi5IVt21mGCfdb/wCO/wAQ7T9pvTvgVB8KdVuPCV5ZNcyeMwZ/7NgkFvLMIGAtTFuLIsfNwDlvu5wpAPqSiog2D/ntS7+cdqAJKKZvyeKC2ODQA+imb/XAoL+vFAD6KZu7Vma1JrCaRet4fjgm1RYZDapcu0cDTbT5YldFdlQtgMVUkDOAaANaivz7/YS/bI8U/taJ49TxT4as/Dk3g25s7ZVtJ5JxK1x5+/cXA+6YeMdQea/QHdzQA+kzTA4I49cUE8A5xzQApbBwcU+vFfjBqPxg8OaVF4t+FVnY6+NISaa+0G5R0udShwrbLK7R9sNwoVvLWSF0lLbS0fDDpvhV8TvCvxi+H+h/ErwRcNdaNr9uJ4S67ZEOSkkUi87ZInVkkAJw6kZPWgD0SiiigAooooAKKKKACiiigAooooA//9D9/KKy9X1rSdA06fWNdvYNOsLUbpri5lWGGNScZd3IVRkgcmuAi+OHwYuJlt7bx94flkc4CpqtozH6ASZoA9ToqGGeK4jSaBxJHINyspyrA9CCOCD2qXIxnPFAC0UZoyPyoAKKKTIoAWikyOuaMj1oAWiiigAooyKTIoAWijNGaACiiigAooozQAUUZFJkUALRSZFGRQAtFN3DNJvoAfRRkdKKACijIozQAUUmRS5FABRRRkUAFNbpTqY/TigD8+f2Lpj4j+Lv7T3jjUPm1Sbx7LojOfvfZNEgWC1X/dCscCov+Ci3xM0Tw38JtD+Eeq6suiD4q6xaaLd3QLbrXRldZtUuAqqxIEIERABJ83j1q78OIo/2fv2uPiJ4Q8SH7F4a+Ns9vr3h68c7IH1mKPytSsGc4X7RJgTxLxuQHBZga1b/AODP7QGu/tpaJ8etbTw1J4C8M6Zd6Rplkt/eNqFul4cy34Q2Ii+0SDCNGJdoj+XeSCxAPy9+F3xO+Fvwx/4KpW+rfCPVre68B/EuCPTM2yvHHFNfQRqsQV1Tk38EZAx91x1r7V/4KEftAfETwz8R/g/+zT8Mtcm8L3XxO1W2h1PVLNjHfQWk15DaxJbTc+WWZ5GdgN3yAZClgZP28P2Nvjx+0t8T/Afjn4R33hzw+/gWNnivdSvbxLuW4M6TR/uoLKZFSBowUJkJYucgYGeu/a8/Y3+Iv7TGlfDz4j+H9Y0zwb8W/AbpOjLJPc6XJL5kc2xZzCswWKVA8TmAnBZSvRgAekeFP2Rtd8F/F/WvFuh/FXxR/wAIPrejSWD6Bc6rf3ckN7KV3XsV7cXMjLJ8pKnZuVnbawGBX5j/APBJP4SaR46ufiJ4l1LXdesbnw7renSxQ6drF3ZW9yQZpMXkMTqtypKAESAgqWB4Yiv17+DOk/tWS6yde/aE1zwzBaQ2f2eLR/DFrcPDJcOysbqe7vf3oZQpVYogEwxLFiAa+Nv2YP2OP2nv2V/ij4xtfAXibwrcfDjxdqUF3LNfw3txqkVtbSyMscdvH5MKyvHIY2ZpmUYDqM5UgHnX7PHijxrbf8FQ/iv8KbrxbrmoeE7DTL6W20y+1W8ureFpDZS/u0mlbb5fmuEI5VTgYHFP/wCCfHi3x3c/tg/tD/DjXPFmta9oHhK8vrLTLXVdTutQWCK11SW3iKm5kkJYRqFLnLEdSa9Y8U/se/tC+FP22NU/am+AmveGo7HxRarbahaa+t2xgDQxQyhIrZR5ozCkqfvo/m+U/KNxy/2Y/wBjD4//AAU+P3xW8eeK9e0fVtC+JialHLqVtczwatFJczvPHdra/ZTCHLt80QnATOVd9uGAPAvj9eWvhP4N/EqTxv8AtEaz4y+N2iTNeWjeGdR1G3stJVJEC20tppv+hwMwVxI86qV45XHPv3hf4meLfi9/wSV13x745vpNT1y78I+ILe5upcGSY2U1zaI7kAZdkiXcx5Y8kkkmvNPhj+wN+1X4B+AfxN/ZqHiXwfH4c8YyXV1DqiJeyancTNFGsUM26JY4YXMSeY2J3QGQIGLK6+5/D39lr9oLwZ+wZrf7KFw3hibxHcQahptpfJqF79iNpqk8k80spNgJFlj851VFRlbCkuvIoA0/+CT2P+GKfC2e9/q3/pZJXz1/wVZh+I3wq1DwF8dfBXizxHpXh65v10vX9O07Wb+ztptv76BljgmRYmkjSZHZNuSEP3iSftv9hT4FfE79m34HQfCP4lyaPdT6ZeXM1rc6PdXE6SxXTmUiRbi2tyjI5IG3fuGCcEc+h/tX/A9P2ivgD4v+E8LQw6jq1sH0+WclY4r+2dZrZnZQzKhkQK5VSQhbAPQgHy9r2n/8L1/bK8E2Hw88Xa9aeDPBXhe11zxAmna1fQ2V818xOjWssccwXzHRWnlYjMsOFdjkV558cLfQrDxH8WLP41/H7U73xBdQ3dz4R8O+GL2/tZ9BhRJJIGubPSeZZBuiXfcoylQTk7iR9R/sP/s06z+zJ8HE8M+NL6DWPGOqTifVbyCSSaMrbxra2dvFJKiOYre1ijRFKgAlgAFxXy18K/2Kf2mfhJ8T/jHL4e8VeGbjwr8XZZzc6vex3c+u2sMr3DDyYQqQ+bi4YMXlZMqj7TgxkA6L/gmd8YPH/wAYf2QvEd98R9ZuNf1Lw/qepaXDe3khmuXtls4LlBLK2Xcq07AM7FtoAzgCvnH/AIJPaVHrP7Inxj0uW4ubPz9SuAJrO4ltLmJhp8RV45oGSRGVgCCrD8sivqb9iv8AZb/aD/Zb+Bvj74Wa23hjVtR1u5n1HSbi21G+EJurm2itjHch9PVo418pXDoJGOSpQABqvf8ABP8A/ZJ+Ln7Lvgjxn8OPipcaFqeleI7gXcE+kXd1LMrvCIJI5I7i0gXaVUMrByc5G3HIAPjX9jrxx8QviN/wTw/aB8R+MPGmvajregHVryw1B9XvBfW0mn6VFeQLHced5qoJUyUDbWBYMCCRXt37Hv7QnjTwZ/wTd8U/HPxdqt34q17w7JrEtvNq1zNePJOjIltHI7uZDEJWUEBhhc4wea5z4Q/sFftY/CD4ZfFj4CeH/GXhYeBvHkF4IbtoryXVGeW3aBYypRIbdJ0CRzv+/MYyY1J5H0B+y1+x1478B/so+Lv2W/jm2lT6drxvBBeaLdzzyeXfphvMSe2gCSQSKrxkM4bOGChfmAPzp0nx98bfiV+ytb/F/wAK6t8Vte+Ot9qj39rd6ZZ65LoX2UXpgezgjt4zpjQCJDIwCE+YDEeBsr94fgT4r8Y+OPg54O8W/EPR7jQvFGpaZbyapZXMD20kN6F2zBoJArxhpAWCEcAjr1r4G+AX7Of7ef7O/gx/gf4L8YeB7rwVb3Fw+n6vfW2oTapYRXUheQxWihIGfczSqkkrqHYguy/LX6e+H7C/0rQ9P03VdRl1i+treGKe9mSOKS6lRArzOkSrGrORkqqhRnAAFAHzx48/aX1HwL4t1DwrF8G/iB4lSwZANR0fSra5sJ96K+YZJLyNmA3bTlBhgw7AnkT+2JqhGB8APil/4I7P/wCT6+zwAKXAoAxdA1Vtc0LT9cksbnSzqFtDcG0vEEd1b+agfypkVnCyJna4DEBgRk9a/Bb9rGfxZB/wVi+HE3gS2s7zX0sNNNnDqE8lvaNJi6yJpYo5XVSM8rGx9q/oC21+Svxo/Y//AGnPGn7aWj/tT+C5fCMVj4WFtDp9hqGpagss8VskihpvJ091jZzIxKozheBubk0AfPHje91i/wD+Cifw4vP257GHQJ7aKxTwbD4fc3Wjy3SXTtA95cTFZ/8Aj5I48pct5YYLGCzfd37RcOlWPxl0vUPjP8bJfCvw4udOMNn4Q0q8udO1TUdSLFTNv04pezwgH7kZIDqvHLA+c65+xj8b/j5+0Z4L+OH7S+u+HrDS/ATW8mn6H4Z+13CTPbTfaFE095HEQGlCmQrGdyLsAU/PVr4h/sm/tC2n7bp/aw+EGseGr63vbCOxey8R/aw1iq2qWr/ZxbRtuDBTICHT5ndCMEtQB47/AMEz/i9411n44fG74N6l4p1fxX4S8NXc02hPrk08t5b28N7JbKD9rAnTzIvLLxuBsZT8isWz8s+HPjfLe/Er4lfDD9u7xv4z+H3j3Ubtk8P6vZ6jqFpo+jE70jeOztJYY2i3bGWRo2jkjyWdD85++P2a/wBjj9oD4E/tF/E/4qavr+ga5pfxHivi95FLc29/DeXEpukuBZG1eAKZyQYvtB2IQQzFdrcx8YP2Tf2xP2gPgxpHwU+Lt34H1e/068jn/wCEzM14dTWBZCxVLRbBEWRlIRisqoyr8wDfNQBnft4/Hn4lfs8fBr4MfBXwD4wu9R1/xlDb2F74ogSS41K6t9OitYpbqDDyuZ7t5Q+Vd3PO1tzBq5TSfFXxk+H/AO1r8OZ/gPovxR1b4YeIRaab4og8V6fr1xbRT3E5ikvjJqcbGIqrpM7oyoNjA4QkV9IftP8A7BTfGf4H/DXwD4H8TnSfFvwjtra30TVL0OVmSGCGFxOY8tG0jW8Um9A2wpgKQ2R658LfD/7bmp67olx8dPE/hLSNH0fL3MHhi1uJ7rWGCGNUuJb9RHbxMT5jeQgcsAAUU4AB8BzA/wDD7G24x/xJzj/wRMP51p3PiDxf4d/4LB6H8P8ATvFGtN4X1O0ur+XSptTu57Hzp9Gupn228srRqnmqHRANiHGwKAoHskn7JX7Qr/8ABQWP9rpW8L/8Iykv2P7AdRvvt/8AZxsjp5mObDyvtGw+d5W/Zu/debj97S63+yX+0Nff8FANM/a7tZPDA8OaW/2NNPfUb4Xr2DWcli8xxYGMTbJWlEW8ruAjMh++QD5+/ajv/iLpP/BTT4a/D3wJ4+17w1pnjOwsbu7ii1Ce4to5ZpL23leG0neS2UmKFdimJoxJ85Q/MDp/BfxH8R/g3/wU/wBb/Z2/4TzxB4t8HaxprTeTr2oS6i6O2nrfBwZflDq6sgZVHyHackZr3D42/smftC/EH9t3wb+1B4al8MR6D4KjsbSCyu9Rvo7y5t4HmkmZ/LsJI43Y3EioA7jAUk8kBp/ZL/aH/wCHgi/tdLN4X/4RnzPsZ0/+0b4X39nfYvsHm4+weV5+z9/5W/bu/deZj97QB86L+0Rq3x7/AGvvin4L+Id341k+HfgEXWkadpHgeHV2d7u3uTb/AGy9fSczE5SRk8w7Cdo2nawb6k/4JteMf2j9U8JeMvA37Qel+I0Hhm9gOiap4lsLq0u76yuRKNhkukV5jCYgzFizL5gUnAAHP6l+x3+0J8Hf2n/FP7RH7KWveG5LTx8Zn1jRfE32uOFZbmVZpXjktEdm/fAyqfkKbmTDr1+8fg/oXxh0nSL69+Nniay17XtUuPOW30q0+y6bpsQUKLe2Mm64lBwXaSdyxZsAKByAfm5/wVt0j4peEPh54e+OXwr8X6/4e/sm7TTNXg0zVby0tpbW43NbzSRQyrGDHMPLLgBn81QxIVQOv8YC1/aN+NvwC8J/CnxX4k0fw7B4VHifxF/Z+u6jAZtElESabbXEkc4L3E83mLJKzCcpuZmOAV+8f2gvhRZfHL4LeMvhPelF/wCEk06a3gkkGUhugN9tKcAnEcyo/AzxXzb+wP8AsneIf2W/hpqVh8Q9Qt9W8Za5cILm4tpZJoIrGzUx2VrE8iRtsjBd8bRgyFegBIB538YNO8MJ8YvF+n/H/wCPGow6drNqq+GPCPhq7vrPUNOBjy9zPbaR/pFywYb4jKjoctkMAAON/wCCRXxp+JHxY+D/AIt0v4k69d+JLnwzqscVpd38zXFyLe5hDeU0rku4R0YruJIDYB2gAa3g39kP9pv4VftV/E74yfDrxN4ZuNE+J0k7SXmsR3dxqemx3E/nhYLeNUjkMBwqK06o6qm7GAK2v2Ef2TPjn+yBpXxE0LWJPD3iCz18i80l4dRvEkN3bqyQxXe6wAjjlVgXlj8xkI4ikz8oB8zf8EhNKtddtP2htEvnmjt9Q1Gyt5Wt5pLaZUmF+jGOaFkkjcAna6MrKcMpBANZf/BPXxf8QviX8Bf2jbDx9408QazLpUKCyu59XvDd2bR2906tbz+cJIjvjUnYyhtuGBGRX1Z/wT8/ZH+OX7Kut+Px8S7vw9qemeM2tbtJdJvLyW4gubV5fkaK4s4VaN0nb5vM3KUA2tuJXy/4OfsKftP/AAL174s+D/AHi/wzB8PviUJV+1XEN1Pq9vFtnWIQwhY4I5SkxR3Z5UHyuEYgxkAq/wDBOH46+M7P9h/4qfGD4i67qPiy98IarrN2kmqXk15KYbHSbS4SBZJWZ1Qvu4BxliRyTn5V8H/E/wCL/wAcf2b/ABZ8WDqvxO1X4zX+qSXGgTeG7XXDoNlDDLEDaRJZg2JV1EgfeGYfICeDu/Qz9iH9jz4k/A/4GePvgB8b/wCxdR0DxlPdymbSL25knaPUbOOyuYZEmtYVQCOMFHV2O5jlRgGuP+Af7MP7cX7LGk6t8KfhT4t8E634HvL2W7sr3XItR/tCyMoVXZbe3AiJbG7yzMyl8nIBYEA+3v2V/GPxJ8f/ALP3gzxT8X9Gu9C8Z3FrJFqlrfWz2VyJ7aaS3MrwOiMhnWMSgbQMOMcYrwD9hCZtM1f9oHwFanGl+HPiRrRsYxwkEN5smMKDsitkge55r6o1HxVp3wO+EZ8S/F3xUdUj8OWStqGrXEMNtJeyqOSkEIVA8rYWOJASSVUbmOT4h+xF8OfFPgv4X614z8f2T6Z4r+KXiDUvF2o2Uv8ArLM6m4MNs2QCDHCiFlP3HZl6g0AfZ1FZWsa3o/h7Tp9Y16+g02wtQDLcXMqwwxgkAF3chVGSByeprl9K+Kfw012G+uND8W6RqMWmQtc3b29/bzLbwIMtLKUchEUclmwB60Ad7RXJeGfHvgjxr9o/4Q3xDp2vC0CGb7Bdw3XleZu2b/KZtu7a2M9cHHQ11mRQAtFc9pnizwxreq6noWjavZ3+o6KyJf21vcRyz2jyglFnjRi0ZYKxUOBnBx0NT634j0Dwxpsms+JdTtdJ0+H/AFlzdzJBCn+9JIVUfiaANqisHQPE/h3xZp66x4V1W01nT3JC3FlPHcQsR1AkjLKT+NVvE3jTwh4Lsl1LxjrtjoNozbBNf3MVrEW/uh5WUZ56ZoA6eis3S9W03W7GHVdGu4b+yuBuingkWWKQZxlXQlSMjHBrRyfSgD//0f3D+JenWOqfDvxPp+pW8d1a3GmXiSRTIHjdTC2QytkEe1fzxf8ABO3QP2cdU/Zf+Ml9+0BZ6DJDb3DCO41NLf7ZEv2It/oskg85ZNwzH5Xzb+gz1/oP+LviLQ/Cvwu8WeIPEd7Dp2m2WmXbzTzuscaDymAyzcZJICjuSAMmvwv/AOCPfgj4EfEDw5450rx/4c8PeIvFem30F3ZxanZWt3eRWJjVGkhE6vII1lwGK/KrMucFhkA9E/4J3fED4z/CP9g/4g/FOfSDr+keGLy61DR7HUrmWySTT7O2WS+a2lEMxKBg3lqq7GkWRdwOSPXPD3/BR340+OPgtpPxd+H37Pmo+IrO41OTTb57W/8AMit5FYbBFHHA9zOChzJIIVijYhd7NkD7I/bG1Twz4J/ZD+Kkd5Lb6RYHwvqenWqHZDEJbm1e2t4I1+UZd3VUUdScD0r5Y/4JB+ItD1L9kxNCsb6K41LRdYvheW6OrTQC4YPC0iZyqyKDtJGDhsH5TgA+gfi5+1yfCnx08P8A7Mfws8NL4y+IuuwPczJc3h07TtNtliaYS3NwIZ3PyKW2JGWxjncyA0fgL+2TF8TPjL4t/Zx+JPhf/hCPiP4RDO9tHeC/sb23TYTLbTmOB+VkSQI0YPltuzkMF+HviNDd/s6f8FVbb46/FQvpnw88c2C21trlwD9htZv7MW18mabbsjcTQYwTwkiuxwWI0f2ffD198cv+Cm/j/wDaf8BxyXHw30a2+xW+sBHS01C6j06308x28hULMNyvJlcqEVWJ+ZdwB75eft7+M/GniX4sab+z/wDDiDxVovwailm1jUNS1Y6ebswGYPHZQx28+8t9nlMbO6hlTkAsoM/ws/bm+JnxV/Zn8cftKaV8NtLtdN8JQTTxWj6/I0lwLBXlvkciwBieOII8QAZZN4BZea/MHwv8TPg14U+LH7Rlhovxgf4G6N461S80e40i+8OXmvzz26mWOa5hltjGtq/myzrHG24xo+DkqGr770G+/Ze+HH/BNX4m+H/gX43g8ReHbbQtZtLnUJ28i4uNX1K2eONZYZFjaN5XeOOOMqMoFA3HLEAj8Pf8FHfjV44+DGlfF34ffs/ah4is7jU5NNv5LW/8yG2kVhsEUUcD3U4KEGSURLFGxVdzNkL9U/Ff9rT/AIRj46aF+zL8K/DieMfiJrdu93MlzeHTtO0y2WNpVlurgQ3DnKqW2RxsSMdC6g/PH/BILxDompfsmLodhfQ3Gp6NrF8Ly2Rw00IuGDxNImcqHXOwkYOGxnBx8jftM+GPDfwe/wCCkNx8Xf2iLLVU+FXjizgji1jT7jULZbWWOwjtdrz6a8UwKTW53RK2fLcPtagD9MP2eP2y9L+Lvif4h/DTx/oA8FeN/hjJP/atmt19utpLa2cxyXNtOI4mZFYDIMYIDIcndx4x4J/bw+L/AMbPAHi34vfAH4Q23ibwx4SvmsjYXGtvF4hvgESTfBY29ncInySAhTMzMQwTcRitL4AQ/sfeHNR+Ivx9/Z18Fa7qQ0fQrl77xFNLqs1vqyAfap7K2OrXJeecfZ42Z/K2jKjzPmIr8pfita/B3wTL4S+Pn/BPzxxqnh7x/wCKdRjt7jwHas019bSSK8skbWybm8pJFCNDKJIn3DYdo20Afsx8af21n+GPjr4afBnQPCA1D4h/Eu2tLmKy1G/GmWOnJdlolFxcGGWQuJUdBGsQZtpHDFVb2L4JfE/4z+MPFvjbwV8YPhyngyfwm1iLTU7O8mvtM1hLxZWdrOWW1tyfI2IJBywZ8OqEYb4S/bEg/ZY+N/jXRvhF+1HLP8NPHOm+HoNTsfE4mSGximuTmaw82QmOXY6ltjgZGfLkVic7H/BKLxf8aPEHgXx1o3j7Vr3xP4L0DUobTwxrV8JM3kK+ak4gabMjQKqQsgZjs3lB0IUA+lP22f2o/FH7Jfw80r4k6P4StfFWmXl+NPukl1B7KaCWaNpIWjRbeZZFIjcOdylTtwGBJX5y+JP/AAUf8efCjw14B+Kni/4J3lr8N/G8doU1U6rB9rL3ECzsYrNEchdm94BNJG0yDJ8o5Al/4LGjH7I9sQcZ8SacD/35ua+NP2ofjJ4J/aG/Yq+BXwG+DNx/wl3xDu30RZNFsVaa8s207TpLW485AP3eJXVQXwrLufO1SQAfrn+0d+138PP2efhjonxDuIJvEtx4teCLw/p9kdsupvcKsilXZSI02MrMzA9QoBYgV5tp/wC2V4r8H/tDeF/2ef2gvA1p4Q1HxzbJc6Nf6ZqzapaGWVnRLScva2rLLvQx7kDKXK4G1tw+Hv8AgpH8CvHXhr4J/s+a/badc+ItH+EVtFp+uR2TSlwogsl8/fEFkSNvsjo0u5dhZDkM1ex/Ci3/AOCevxp+J/gfWvhL4d8SePPFmnXUF/BdyXviG5Tw+bZhcpNey6jdi3REmVV2K0u9yAEdTyAej+MP28fiD4M/a2k/Zaf4Wrrt5cJJLp0+m6oHmvEe1kubbKTwQxwkhcTlpNsQDspkCjd6B8P/ANsrWovhL8Rviz+0h8Pr34U2vw/1A2TWkrPeSXf3Y0Fu7RQpMXmYRK8ZMbEhtwXkfn94s+IPgm2/4LR+HdVuNatI7KwhOkz3DTIIkv5NHuLdIGfoJDNIkO3O7zDt68V96f8ABS74XeNfir+yZ4h0L4f2c2p6np1zZ6ibG2UvNcwW0n7yOONQS7KG8wL1JX5cnAIBzWl/tl/tAatpfw2+INl8Cnu/A/xKvoLW3ks9Umu9SsLe4cBL29ghsWjjhaM+YuJCNo+ZlLCvTfHP7Wt8/wC0BJ+y/wDBDwvB4w8cWNg2oarPf37abpmlxbI3UTSxwXMsjkSx/LHHgGRBuzv2eGfsZftsfBvWfgx8LPhBpKapqXjzS7DTtAvNFstMuZZrVrNFtZLqeUotuluqr5zuZchcjaZAVr4l+Ifh7wN8CP8Agov448XftVW+r2Xw88f288ml65ZXWp2cSyzfZ5UDT6XJHM6xeU8DxZbBKSFdoVgAfqB+z5+2Jc/tDeHPiFpOk+E49C+Jnw4nms9Q0G+v82zXCNJGuy9jhLeW0kTozeSShAyCCpPzJ8L/APgph8UPi78IfHHxK8EfBF9Sm8Ch7nUANZjis7exSAzNI0ssSySzDY5EMUR+RSWZCUD+2/s0WH7IfhCy+Inxq+Bvh3UvDmgi1zqPivV59SWy1RV3zzvC2q3DyuIXH72YxIGZsK7/ADGvzL/4J2+PvB3hz9jP9pjT/EGsWtjcw6ddTmKeZI3MV1YSWsO1WILeZOyxKBnLsij5mUEA/Q7X/wBsTWvjL+wF4x/aA+FugSaZrCWGpWF5A1+Y5NJkSJkkuYbhY1aZoldJYgFjYlgCVwa57/gm98U/j74g/Z88Oap440T+3vCqLq88niS61uW91edobiYiP7FLAzHDAxqTcH5QCBghR8qfsQxnx9/wTG+N3w48Juup+J4G1mRtOhO+623FlE8BEQy374xSLFx87oyrkg49V/4J0/tLeGtA/ZFn+FnhOyvNe+Jng618Q366LHYXbq8sRnvLdZriOLyYxO7LCu6QM0hCqN2MgHt3jb9vT4weHPgrc/tK6f8ABNofh3ZXy2rx6xq7WGtvE84tluDZraSxxRmZggzM7E8hdvzV9H+Lv2vPhv4P/Zbs/wBqu+trmXQNRsLS7tbJdn2uWe8ZUjtvvFA6yMRIQxCqjMM4xX4d+NPjtqn7Rv7GvxI1v4l+OvEfiL4nLexNJ4cs4prLQtK0+3vYHEtzBbxJbsgUNta4kd/N8sBfMG4/TGu+Lfhpqf8AwRu05tat28RR2FnBpojsZ1WSx1dbz9y07gSeWIiVd1YZkjYKNvmKQAez+Pf+Cjfxn+G3hv4a+OPFnwe0y00D4tRxy6M58RSs9vHJ5ZU3pGnnZuSZJBsVvkz/ABAgfUXxm/a6l+Ddx8Ovh/feEv7Z+LPxGe2htfDtrfBbe3klIjleW/aEgQRykqJPJy4BbaFU1+TPg/xJ+yb8Rbb4OeF/jp+01F4h8P8AwujgXTtFi8L3+jqzr5IWK81CXzEeOMRrFu2x5jBIZSzNXrf/AAUe8K/8I/8AtVfCH9pPxbp19rXwmisbPT9Tu9KmuI2tUS5nkdxcWTxyRb47kSRFZF8woyhutAH338Jv2v8AUtf/AGidb/ZW+MHhWHwl47022W9s5bC/bUNN1GExLOfJlkgtpVcRsTtaMg7HyVK4P3IvHHQCvzD/AGdNC/Yc8d/HDTviH+z14c1vxR4h0W0uJG8VzXWuT2Nh5kTWy20surXIEkksUsgSJI5NgVidpUEfp6BzQB5z8UPiZpvwq8Nf8JRqukavrVv56QeRounzandbpASG8iBWfYMctjA4z1r53/4bd8DD/mnnxE78f8Ifqfb/ALZV9EfE/SPiZrXhsWXwn8Q2XhjXBOjm7v7BtRh8kBt6eSssJ3ElSG38YPHNfO5+HH7cY/5rN4Y/8JJ//lhQB678JfjzoPxgvNRs9G8N+JNCbTY0kdte0W60pJBIxUCJrhFDsNpLBckd69wJPSvDvhH4Z+Peg3moyfGTxvpPiy2ljjFpHp2jtpjQyAnezs1xN5gYbQBhcYrE/ae/Zg8A/tYeArD4d/ES/wBT07TtO1OLVY5NKlhhnM8MM0CqzTwzrsKzsSAoOQOcZBAPZPF3gvwl4+0Kfwz420e01zSbnHmW15Ck8TMv3W2uDhlPKsMMpwQQQDW7YWNrptnb6fZJ5dvaxrFGuS21EG1Rkkk4Axyc1+K/if8A4JQfsReC9V0bQ/EvxD8V2WqeIrhbbTrQ6hp73N1IxCnyoU05pGVdw8xwu2MHc5Vea7Uf8EWf2W2P/I0+MvwvtO/+V1AH7AYFGBX4/wD/AA5Y/ZZbGPFXjIg+l9p3/wAr6f8A8OVv2Wf+hq8Z/wDgdp3/AMrqAP1829R60bRX5B/8OVv2Wf8AoafGf/gdp3/yuo/4crfss/8AQ0+M/wDwO07/AOV1AH6+7RSbRX5B/wDDlb9ln/oafGf/AIHad/8AK6j/AIcrfss/9DT4z/8AA7Tv/ldQB+vpAPNJtFfkF/w5X/ZYxkeKvGf/AIHad/8AK6m/8OWP2W+f+Kq8Z8f9P2nf/K6gD9fyooxzmvyB/wCHK/7LX/Q1eM//AAO07/5XU7/hyt+yz/0NXjP/AMDtO/8AldQB+vm0Uu0V+QX/AA5W/ZZ/6Gnxn/4Had/8rqP+HK37LPbxT4z/APA7Tv8A5XUAfr7tGc0YB5r8gT/wRY/ZZAJPirxngf8AT9p3/wAr653Sv+CQn7GGu3cmn6J8Q/E2oXUSeY8Vvq2lSyKmQNxVLAkDJAyaAP2g2ikCgcV+Qf8Aw5W/ZZx/yNXjP/wO07/5XUf8OVv2WP8AoafGf/gdp3/yuoA/X3aKTaK/IP8A4crfss/9DT4z/wDA7Tv/AJXUf8OVv2WP+hp8Z/8Agdp3/wArqAP1+or8gv8Ahyr+yz/0NPjP/wADtO/+V1H/AA5V/ZZ/6Gnxn/4Had/8rqAP19pMCvyC/wCHK37LP/Q0+M//AAO07/5XUxf+CLH7LJ6eKvGf/gdp3/yuoA/YAgGkx3PWvyA/4cs/stYyPFXjIj/r+07/AOV1L/w5X/Za/wChp8Z/+B2nf/K6gD9fto6UbRX5Bj/git+yz/0NXjP/AMDtO/8AldR/w5W/ZYHXxT4z/wDA7Tv/AJXUAfr5tFG0V+Qf/Dlb9lj/AKGnxn/4Had/8rqD/wAEVv2Wf+hp8Z/+B2nf/K6gD9fcUm0V+MN//wAEhf2L9Kvo9L1T4h+JrK9lKqkE2r6VHKxc4UKjWAYkngYHNdB/w5Y/ZZzj/hKvGf8A4Had/wDK6gD9f8DrRgV+QX/Dlb9lnv4q8Z/+B2nf/K6j/hyt+yz/ANDT4z/8DtO/+V1AH697R+dKFr8g/wDhyt+yx/0NXjP/AMDtO/8AldR/w5W/ZZ/6Gnxn/wCB2nf/ACuoA/XzaKNg71+Qf/Dlb9lj/oavGf8A4Had/wDK6j/hyt+yx/0NXjP/AMDtO/8AldQB+vu0U3YBX5Cf8OVv2Wf+hp8Z/wDgdp3/AMrqP+HK37LI/wCZp8Z/+B2nf/K6gD9fNoP40bFr8YLH/gkL+xfqmpPo2mfEPxNd38YZmt4dX0mSZQuNxKLYFgBkZ44ro/8Ahyt+yz/0NPjP/wADtO/+V1AH6+4ppX0r8hP+HK37LH/Q0+M//A7Tv/ldR/w5W/ZZ/wChp8Z/+B2nf/K6gD9S9f8AAHgrxTrui+JPE+i2uq6l4deSTTprqMSm0kk27pIg+VWT5RhwAyjIUgE56xM8jGB+tfix41/4JKfsWfDjRD4m8b+PfF2j6QkiRSXc99p4giL9GlkGmlY045dyEHGTXqvwz/4JNfsy+B/GfhP4p+FPFPijUbrw9f2Os2HmX2nzWs72kqXMO7y7FC8TlRna6kqeGHBoA/TfxR4a0Lxl4c1Hwp4ms49R0rV4JLW6t5RuSWGVSrqw9wfwPI5r+cT9l+T4q/A74nfGr9gzwJo8N14u8XX4sLHV7tIxDYafEkwuNQnVuZVNlJHNBGBgueQd2D+uv7ZH7Y/hn9nHSLDwdoV9YT/EvxY8VtpNneTCO1s1uJDENQv3z+7tYmyeSDIVIGFDun5NftXeDrH9g79o74H/ABp07xRN4o8YXQmvfFcs04N7qEnn7by68ot+5ju4biW3iUYRVhAGSHJAP3X/AGff2ffh5+zb8N9P+HHw8shDDAqtd3jKPtV/dfx3FwwGWdj0B+VFARQFUCp/jN8MPFnxX0uy8MaJ491LwJpEpmGqvo0cSaldwsoWOOC8lDm1wclnSNnYfKCo6+l+FfFHh/xn4d07xb4UvotS0fV4Irm0uYWzHLDKoZGX0yCODgjoQDXL/E74u/DL4OaEniX4peJbLwzpsjmOOW9lEYlkCM5jiXlpH2qSEQFjjgHpQB+Nn/BE5JFX42iaRpZVvNFDOxLMxH27lj355zXX/FvWh8fP+Cq3g/4E+NkXUvBPw+sjeNpMwElpcXzaa98JponyjtmaIYYH5UxjlifBP+CPfxo+F3gbXfib4a8beJLLQNS8W3mjnS4L6VYDdsGu0McZb5C4aWMbN2SWAANe7fHbRpf2cf8Agpp4U/ag8aRyWvw48YWf2O91jyna2sbsac9gsc7IDt/1cMm5sDazEf6tsAFP4QamfgT/AMFZfGXwV8EKNN8GePLQ3J0qEeXaR3X9mR6h50USYRWDxyqpAGEkK+mJvg1d6f8AtQ/8FNPi1J8ULODxFoPwwsb7S9G0y9RbqztpLa6hsnkSGUMm9ysrscA72H90Y3/gN4Uuf2gf+Cknjf8Aas8LxPcfDnwxaDTtM1XyysGo3n2COxkFuzY8xBmdt65XGzn5hXJeBJNO/Y0/4KO/E/xJ8Ypv+Ee8D/FK1vrzStbuUZNPe6u7qC9eJ5gCiMj+cjBiDu2k8OpoA1P+CcviG++Hn7XP7QP7MGmzMvhHR77UdR0qyZmaOzW01AW2IgxO0PFPGH/vFFP1/bWvx4/4J4fDnX/FH7Qnxw/a3vdOuNO8NeN9Rvrbw69zC0El7ZT3zXD3Co+GCARwjOCGYsM/IRX7D5FAH//S/foqCMU3buGGqSigBhXuOppu3jA4A/pUtFAFaSCOZWjlQOjDBU8g9uc9fxokjZozHGzIcYDDGR6EZBH5j8Ks0UAfj/8AB/8AZJ+M3wa03xP4Z8cfCnwL8c77XNWutSj8Ua3f+XezC5VVIvUu9Pu5CSVLsschGWK5bJevpP8AYd/ZIvv2XNC8YXviK9sZ/EPjrUVv7u10lJV0ywiiMnkWtqZgrukfmP8AOyIcFV24XcfuzAowKAItvqOPz/DpSSwRzI0cqB0bhgwBBHcEHI5qeigCEJjt7fWmeREJTMEHmEBS2PmIHQZ64qzRQBXkt4p0aK4QSI3VWG4H8DTtuFCjjHAqaigD4A/4KG/AD4tftNfCDSfhf8K4NND/ANqxajd3OpXb2wiS2jdEjjSOGUu0hlJJJUKE6EsMfTPwC8MeKfBvwc8HeFvHVjaWPiDQ9Js9MuxZTm5hc2ES26SLK0cTEOqB9pX5CxXLY3H2TaKNozmgBu329Kijt4oUKxIEBySFAAJPU/U96s0UARFPQZ7/AI+1Gw8+p/z/AJ/KpaKAIFijVmdECs5yzAYJ4A5/AYokgjmRo5kEiN1VgCDip6KAItgznHNJtOOenpU1FAEe3rxUaQRoWMahC53NtGMseMnjk8dasUUAV0hSPcYkCFzk7QBknqT7/wD1q+Nv25v2ffiN+0l8FU+H/wANdet9H1CLUra+nt71pY7HUoIFfNrcPCGkCF2SQDDDci5AOHX7RpNooA/IT4xfsc/ED9oXwbZ/DiP4L/D74QBrq2kvfEOnzpe38cMRzILKG20+1Pz9Css6gjKns9fqd4M8J6f4G8GaF4I0su1j4esLXT7cucsYrSJYkLHnJ2qMnHWus2jOaNooAjihjiRYolCIg2qAMAAcAD2FTUUUAJjnNLRRQAmAKq313b2NnPe3biKC2RpJHPRUQZJP0FWzXNeL9Hl8ReE9a0CBtkmp2VzbKx6K08TID+BNAHxf+wzFL8T/AAprH7WXiyMz+J/ihf3jWrSjdJpuhWdzJb2Wnw/3Ix5ZlfbjzHfc+WAI+BP2rv2RbL9n/wDZH+IfxO8Qazdax8QrvXoru31KG/vUitLS81GNY7eKHzBHkROwYlCcscHABH6Ef8E4NTF3+xz4D02ZTBe6CL/S72BvvwXNpfTxvG47NjDY9GFcP/wVcuYIv2J/F6SOqNPeaQkYJALt9uhbAB6napOB2BPY0AfPv7M37H+n/EH4N/s6/Hjwbq1zo/ijSdQstY1+Se+vZYdWs4Lh2kiMTSPGsuUUKQqqQWDZ4x9xeL/2wvCuleOfFfwy+HvhDxB8RPE/gi3Fzrdvo9vBHFYIVDhZJ72e3WSR1OUSDzWbBABIIGb/AME9poLn9jL4VvbusippjKdpBAZbiUMCexByCPXivkq1/a+1Px7+078VvhnB4t0P4IeGvAJuI7vU5rO1n1vW57GY27MjXn7koDuZFEMj7WUAHedoB9tfA79rfwF+0V8IdY+LXwq0rVNT/sOSeC40aSO3g1L7TAgkES75hbkyIwKN5205wSrBlHzDo3/BVz4LeIPhvrvxI0fwV4tvbbw9cRxXlvb2EUzW0UoGy5uZ0mMEELOfLUySB2fO1WCsR87/APBGC5gsPhj8X4b5xAbLUbWSZJfk8tfs8uS6nAUfKQenT2q//wAEUPsN58KviZYzCOaT+2LNpYyAx8t7YhSynsSrYJ64PpQB96+O/wBt74WeA9H+Grvpera14j+K9nZXmh+H9Oghk1B4r5EeMymSWOCMAuFJMvXJUEKxHK/DT9vPQPiV8aLr4BQfDTxVpfi/SZCmrRXK6cYtPjDrG00rR3jM8Ss6ZeJGGGDDKkGvkD9vUeF9D/bb+BXibwn4m0jw1490awNw/wDwklwNP8PJpdnLcS2/n3Yy0TTS+fCirG2cqflwN3sn7JPwu0TW/wBqH4g/tR+JviX4P8VeMvFFp9lg0bwhqianbWFofITfJKRHK7hYI0B8tV5ZiSWAUA+i/F/7Z3gvSvHvin4a/DvwvrfxK13wPatd68ugpatb6cked8Uk11cQK864P7mLe5IKgb1cLx3jX9tDwp4r/Y48V/tFfAyLUNVaCxvraBUgjF1pWoLC217yN2KKsBZJHwXDKV27gwr8uv2RYNI8JftHfH34f/Gf4tar8IPEt3qhm8yC602wi1Py7m7aXzJtStbhCcTJJEEZPMR2bDAAj660j4I/Ar4cfsJfH/S/2Y/E2reNtC1Gy1MzXV60c8T3djbAS/Y5YbW3jni2DBkjMiFlIVuCAAesf8E9f2mfir8Z/hR4Zj+JfhvXtZ1C+bUTN4se3sItJlEFxKI4wIZUlDKqiLi3wWXk4y1fQ37Un7VHhr9k/wAL6X408b+GNZ1vQ9RuTZtdaSltIlrOVLxpMJ54WUSBW2soZcqQxBKhvkv/AIJUfFj4at+y74M+GZ8T6cPFq3usINJNzGt82LiW6LLblvMKiJwxcLtxnnivsr9rX4Pp8eP2dfHPwyWES3+o6e8ungjH+n2pE9qM9QDKiqx/uk9jQB5/4+/bP8NfDv4KeBPjnqvgbxFeaN8QJLWOztrNbCa9ibUFMlkJYheYJuY/mQRNIRwsm1vlrofFv7WPhrw3440n4RaZ4U17xH8SdU0tNWfw7p8NsZ7G3dck3d3PcRWKBWBRilw43DjduTd+YP8AwTx8ReJv2nrP4WeBfFFpL/wjH7Ogur+5kmGY73Vp5JIdGjwehsbcysOhBVOOa+kPi/8AtX63/wANsN+zZ4f1TQvhba6dpySal4w1W2hmvJ1e2S8FtaNcPHAqkMq5k3/OrEL8oDAH1b+zZ+138O/2mrjxRoXhvTNU8P8AiLwZcLbatpOsQxxXELMzpkGKSRGCyRujDcGVlO5QCpP5tfsIaLpfh/8A4KUftFaLodnFY2FpFqywwQoI4ol/tS3IVEUBVUZ4AAAHSsn/AIJlajb3f7an7Q0yay2tjUWvruK+kRIXv4pNULi68uNUjHmh1f5EC/ONoAIFdB+xJfWcv/BT79o9YpkZpU1kKAwOfL1W2DY9cHr6UAfe3wW/bY8H/G74keOfhDoHg3X9L8V+AYLqS+s78WKB5rSYWz28UsV3JGZDKdqsWEZHO/HNbH7L37ZPgD9qvUfGWl+C9C1jRLrwRJaxX0erR28ZL3RnUKnkTzfcNu4bOO2M84/Mn9mHx94J+EH/AAUm/aOb4pa5Z+FYtROrSwTalOlpHJuv4rlFRpCoZnhbzFUZLKCVBAre/wCCQmp2+p/EX9pCazlML39/plxAsilXCmfUyGMbYb5dyhgQCDweaAPvXxT+3T4AsNY8c6X4A8K6/wDESH4ZxGfxJfaHHZtZ6eihzIokubqEzSoI3ykKtja2T8pr6M+DPxk8CfHj4e6Z8TvhzfG+0TVQ4QyJ5c0UkbFJIpUOSrowII6HggspDH+fj9gvSPCFkPi98H/jp8X9W+EevpfSRalYi+0vToNQTbLBdCSTUbSffIjAhtkg+VsgYya/ZH9iD4X/AAS+FXwaudI/Z+8Q6p4p8I6nq95eJfamQxedVjtZTbOlvbK9vmAFHRWVjuZXOcAA+yN2elOPFeCfFf4BWfxZ1ey1e58deMPCjWUBg8jw5rk+lQSjcW3yxxDDvzjce3FeVn9i3Su3xk+KX/hYXn+FAH2YWJBwK/Bn/gtNE2g3nwo8T6JPPpuo6kdVs7qaCaSLzoLf7M8KOEYA7GmkwcZG49sV+1Xw18AQfDTwrb+FINd1jxEkEkkn2zXb59RvnMjbtr3EnzMq5wo7CvxW/wCC301u9l8GrRnUy/aNbYpnnbiyGSOuKAOf/bx8EeAfg7qfw6b9jHxFqOn/ABM1TVGt/wCytB1q81Ce5g8vcrvC1zMVxIFUDAWRXbIYLX60fEX9ozw78BfDngPTfivFeal428ZiGxs9K0a2NzcX+pqkS3CQcpCgEkowZJUHIx3x+T//AAUe0v8AZ++EGneAvFn7KsmleE/iuutrHAvhCSK3uHsngff5kFkQp3TCBV3Id25l5BYV9R/tWftZePPg83wH8C3ml6NpHjnx7FZ/2j4h12382y0CSY28N3LEm6MBkkkLOTIqoiDcpDAqAfR/w2/bU8FeOfjpffs2+KPCuueBPH9rA1wllq8drJFcRrEJ8RT2lxOjHyiZOuzAIDFhirni39sPwpp3xj1P4AfDjwprXxH8daHZ/btStNIW0hgsoR5ZIluL+5toy+Jo/lj38sFJDZUflT4e1dLv/grv4Bu5fHafEHFlPA+sRw2sFvLIdJvUMUP2VRCyo+U4LtvDIWLqQO//AGlfhd4K8SfEb4q/tVfspfFtvBHxM+Gj3UXifTZ2NqlxLp0SiXyi+CRMsYXDJLBPIMZU5NAH6d/Dr9qz4beOPglq3x41qG98E+HdAmu7bUU1yOOG4tZbJxFIjJFJLuYyHYij52bChNxAPjI/4KI/DCwtfB/ivxh4Q8S+FfA3j2c22j+JNRt7RbGZwSA0scN1LcQRN95HliG5QWA2gsPif40fE74x/tR/8Eon+I/iDT2/tmC/ik1V7eHykvrDT7wo10kSnAQMEeTHyho3YKqDAx/DPgz9kX44/sofDfT/AIxftDeIZbPS7PS408Mw3+lPdWeqQwfY/s1pYRaa99IAXZIlw7NGQxLfeoA6D/gp3o2kWf7V37NXiC0tIYr/AFLVY47m4RAskyW9/ZeUHcctsDttznGeK9l/4KCeIfganxp+CWk/GzRvGEd5p2otc+Hr3Q5dNTT7i5mubQTJci43z4haKHO1I/lc7SxJ2+G/8FL49M8O/H/9lXRReNLHo95ErSXLgzeTHe2KCSY4UZYIcnaBkH8Nb/grZf2Vv8Zf2ahNOiG21O9llyQNkZu9Nwzeg+VuTxwfQ0Affn7Vf7aXhD9ka58PDx94V1jUtP8AEfmLb32nfZGt1miI8yGQSzxyKyqytnZtIPBJDAcsf+ChPwntvjP4a+DeueG/Euht4zeGPRdY1HTvsmn35uX8uCSITOtx5Ur4RHaEckEqqncPi3/gttc2yeBfhZbu675NU1GQITyyLDEGIHXHzAH61B/wU81LTI/2n/2YbqO4i2wajHOzhhhIDqNmUkJHG07WwenBoA/UD40/tO/D34K+JPDXgK/t77xJ408YS+XpOgaOkM1/cDJXzW86WGOKHcCDJI6r8rYyEcjL+CP7V3gb40+NvFHwrGl6l4T8eeDT/wATPQ9XjhW4WPKqZYnt5ZoZYwzoNyv/ABK2NrKT+VX7Xav4N/4Kf/D/AMZ/EfxVqPgTwnqukpbWPiKxaCJrMG2ubdgktzDPAqieUCYunyxy7ztBDD7N+AXwe/Zltf2p9W+KfgD4t618TPiTJpLPqkjXljqFl9lkEVsn2ibTrGKFJcRp5cfmq7BS2xgCaAPp/wDaY/aP0P8AZf8AAK/EzxZ4c1bXtCS4S2uZdJS3ka0aU7YnmWeeE7HfCBl3YYgHGRnzTxD+254Q8Nfs3aL+1BeeC/EEvhXWZo0SCH+z3vYorhzFBPJGLzZtkkAUKrtIpYblHOPfPjt8LbL40/B3xh8LdR2hPEmm3FrG7dI7grut5T/1ymCOPpX4e/sB6h4w+PeneBf2WvGGnzReHfghr1/4h1zzgNkpt3H9lWD9clb2W4kkjYFSkKgZxwAfrt45/av8IfD3/hBtE8QeHNdbxt8RY9+k+Fre1hl1RioBdZX88WcRj3DeWuAFGScBWxhfBf8AbK8E/F74t+JPgNqHhzWfBXj7wzE09xpmrxwYeNSm5oprWaZGIEiNyQCrBkLgMR4P+2L+1Trfwy/aH+GfwL8KpovhnUfFMS3E/jDXbYXEelWtzNJbkWyu8SiU+U24u+zDopABLD4z+AOq219/wV01m9XxgPG63mmzRR615dvAl+8OlQJII1tUSFliaJo1KA5EW4sxyxAO2+AGjaV4e/4LA/FXSNBtItPs00q6kWCBBHGrzQ2MrkKvA3SMXbjliTX6FfDT9tLwT8Sv2gfE37Ndt4V13RvFvhWC6uLn7aln9ndbZ41/dvDcykmQSqyZABHUg1+f3wWv7Nv+CyXxTKToTNplxCg3D5pEtLHeg56rsbI68H0NYngPxl4T+E//AAV6+Keo/EjWLTwzp+p6XLFBdajOlrbs01tYzx5llKoNyRsRkjJ4BzgEA/R/9nv9szwN+0P8RvGPwt0Tw3rXh7X/AAKZE1OHVUtQqywztbSRq9tczhmSRSCcYPYmm+L/ANs7wXpXjzxT8Nfh54Y1z4la74GtWvNeXQI7QwackZO+KSa6uYFecYI8mLe5IZQN6sq/nV/wTZ8Q6T4j/bh/aT1vSZd1prd5qV9a7lKNJbzatJIj7Gww+WRewxkZrx79kOLSPCX7R3x98A/Gn4t6r8IPEl1qhn8yC602wi1IpcXTSmSbUrW4UnE0ckIjZd6OzYYAEAH7m/Ar46fDP9pn4bQ/EX4czPe6PePJa3FvdRBJoJ0A8y3uI8soYKwJALKVYEEqefnn9m++Pwp/aG+KX7Kdu7f8I1psFn4r8L25ORZ2GosUvbSIH7sEN3zCo+6HI6AVofsK/CD9n74TeBfE1v8As4+K9T8XeHL/AFZ1nvL+SOaA3dtGsbm0mitreOaLGFMkZkQlcB8ggc94KjfxN/wUs+IniLT/AJrTwb4B0vQbxhyFu9QvBqESk9M+UpOKAP0FCnvRs6ZJPSpaKAIgmMADpShRnJGcZqSigCJkHpmkMaurJIu5WGCDyDx+vpU1FAEKx7QFA4HGOw9h7e1JJEkoxKgcdcEZ6fWp6KAIwvzf55p+BS0UAf/T/fyiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKRhkYpaKAPlrR/g14u+FfxW1jxt8IZ7J/C/jq/S+8R6DfvJAsN8+1J9T06aKOTEsiAGaCSPbMygiWI5z7N44+Fvw0+Jtva23xI8I6T4sisWZ7dNWsLe+WFnADGMTo4UsAASuMgDPFd33NSUAcX4L+HngL4b6VJoXw78Nab4W02aZrh7XSrOGygeZgqtIY4FRd5VVBYjJCgdhWNe/Bj4R6l41i+JGoeCdEufFcJVk1aTTrd79WTGxhcMhk3KAArBsqOmOlem0UAeQab8AfgXoravJo3w58OWD+IbeW01I2+j2URvbachpYbgrEPNjkYAusmQxGSM4rQ8FfBf4P8Aw1vZ9V+HPgTQvCt7dR+TNNpOl2tjK8W4Nsd4I0LLuAOCSMgH0r0+igD8iv2q/wBnrxh8Qf2tNG+JHxS+GF18WfhBZaC2n21lok9vBqFneOS7PPG9xaTTjfuKgS7FDAj5wVfj/gt+xVNa/tieFvjn8M/hrqHwZ+HPhWykMtpql+k1/ql7NHcRHy7eO5vGgj2yRiQSSKGVThdzNj9oD0P1pT0/GgDzDxp8Efg58SNSg1n4heBNC8TahagLFc6lpttdzIqkkKryxswUE5AzjPbNd9Y6LpOl6XFomm2UNpp1vEIIraGJY4Y4gMBFjUbQgHG0DGO1aveigDzPwZ8G/hL8ONQvNX+HvgrRfDN9qHFxPpunW9pLKuc7XeFEZhkZwTjPPWvR9p61LRQBxPg/4deA/h9FfQeBPDth4ei1S4e7uk0+1jtlnuX4aWQRqu9z3Zsnisjxb8GfhH491yy8TeOfBOi+IdW05Qtteahp1vd3ESqdwCSSozAAnIGeDyOea9NooA8zsPg38JNK8ZSfEbS/BGh2niyVpHfWIdMtY9QZ5VKyMbpYxKS6kqxLcg4J9crwx+z58BfBOvW/ijwZ8NfDWgazZl/IvtP0eytbqLzUaOTZNFErrvRirYIypIPBNew0UAeZ6t8G/hLr3jC3+IOueC9F1DxPZ7TDqlxp1vLexmPlCs7IZAU/hO7K9iKg8MfBH4OeCvElz4z8H+BtE0XxBe+b5+o2WnW8F5L57b5d86Ish3tgsC3J57V6nRQB5P4s+BPwT8e65H4n8c+ANA8Q6xEFAvL/AEu1urjCDCjzZY2fC9hnjtXptnZWthaw2NlClvb26LHFHGoVERAFVVUABQAMADgCrdFABRRRQAhryvxn8Dvgx8R9UTXPiH4B8P8AijU44Vt1utU0q0vZ1hQlljEk8bsEDMxC5xkk9zn1WigDyTwt8Bvgf4G1FNY8E/Dvw74fv4s7LjT9Js7SZc8HEkUSsCR15re8d/C/4c/FCwt9M+JHhbTPFVpaSebDDqlnDeRxyf3kWZHCk9CR1HByK72igDyS/wDgL8D9T1HStY1L4eeHbu/0OOGHT7ibSLOSazit3MkMdu7RFokjclkCEBWJK4JzWf4l/Zx/Z+8ZatJr/iz4Z+GtY1OaQzSXV5pFnPNI7ElmeR4izliSTknJ5PPNe10UAZlppWnadp8WkadaQ21jBGIY7eONUhSNRtVFRQFCgAAADGBjpXmnhr4BfA7wZrx8VeEPh54e0TWSSRe2WlWlvcqWBDESxxqwyCQcHnvXr1FAHkvi/wCA/wAEPiFrB8RePfh54d8Saq0axG71PSbS8uDGgO1DLNE77VycDOBUXif9n/4FeNry31Hxn8OPDevXdpBHawzX+j2d1JHbxcRxI0sTFY0H3VBAA6CvX6KAPJvFPwI+CXjhrB/Gnw98O6+dKtks7M6hpNpdm2tkyUhhMsTeXGueEXCjsKztZ/Zw/Z68RNZv4g+F/hfUzp9tDZ2xutFsZzBa267YoIt8J2RRrwiLhV6ACvaqKAOE8VfDL4deOtCg8L+NfC2l69o1tt8qyv7KC5tovLXahSKRWVSo4UqMgdKteDPAHgb4daWdD8AeHdP8NaaW3m2020is4S+ANxSFVUtgAZxmuxooAjZcjpn2rj/DHw78CeCtR1rV/B/hzT9FvfEdx9r1OeztooJb24LM3mzuigyPl2O5iTlie5rtKKAPO/HHwm+GHxOWzT4k+ENI8VLpzF7YapYwXnks3Ux+cjbd2BnHXvms6b4HfBm48R6f4xuPAPh+XX9JFuLPUW0u0a8tvsgAgEM5i8xBEABHtI2AYXHFeq0UAePaX+z58BtF8Sp4z0b4beGrDxBFO1ymowaPZRXizuSWlE6RCQOxJy27Jyeeta3if4O/Cbxt4hs/F3jHwVo2ua5poVba+vtPt7m6hCEsojllRmUKx3KAeDyOa9LooA8r0n4IfBvQfGc3xG0XwLodj4ruZZp5NWg022jv3luc+c5uAgl3Sbm3ndlsnOc0vjT4I/Bz4kajb6x8Q/AuheJtQtAFiudS022u5kRTkKskyMwUE5xnGe1ep0UAec+KrHxdoPg1NG+Dml6TFqEIS3tIb93tNPtIQMb/AC7aJmYRj7sS+WG6b0HNcl8C/grY/Bnw/qUdxqUviHxP4mvpdW17WbiMRS6hqE/3nEakrFDGoEcMKnbGigcnJPuP8f4UrdKAFooHSigAooooAKKKKACiiigAooooA//Z"
