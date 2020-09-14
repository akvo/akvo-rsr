# -*- coding: utf-8 -*-

"""Akvo RSR is covered by the GNU Affero General Public License.

See more details in the license.txt file located at the root folder of the
Akvo RSR module. For additional details on the GNU license please
see < http://www.gnu.org/licenses/agpl.html >.
"""

import os

from akvo.rsr.models import Project, IndicatorPeriod
from akvo.rsr.decorators import with_download_indicator
from datetime import datetime
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from docx import Document
from docx.shared import Mm

from . import utils
from .docx_utils import markdown_to_docx, set_repeat_table_header, change_orientation


EUTF_ORG_ID = 3394


class EUTFProjectProxy(utils.ProjectProxy):
    def __init__(self, project, results={}):
        super().__init__(project, results)
        self._custom_fields = None

    @property
    def contact_person(self):
        contact = self.contacts.first()
        return contact.person_name if contact else ''

    @property
    def actual_start_year(self):
        return str(self.date_start_actual.year) if self.date_start_actual else ''

    @property
    def planned_end_year(self):
        return str(self.date_end_planned.year) if self.date_end_planned else ''

    @property
    def cofunding_partners(self):
        return self.funding_partnerships().exclude(organisation__id=EUTF_ORG_ID)

    @property
    def eutf_funding_amount(self):
        eutf = self.funding_partnerships().filter(organisation__id=EUTF_ORG_ID).first()
        return '€{:,}'.format(eutf.funding_amount) if eutf else ''

    @property
    def cf_relationship_with_beneficiaries(self):
        return self.get_custom_field('Relationship with beneficiaries')

    @property
    def cf_synergies_with_other_actions(self):
        return self.get_custom_field('Synergies with other actions')

    @property
    def cf_cooperation_with_contracting_authority(self):
        return self.get_custom_field('Cooperation with contracting authority')

    @property
    def cf_eu_visibility(self):
        return self.get_custom_field('EU visibility')

    @property
    def cf_additional_comments(self):
        return self.get_custom_field('Additional comments')

    @property
    def cf_sustainability(self):
        return self.get_custom_field('Sustainability')

    @property
    def cf_executive_summary_of_the_action(self):
        return self.get_custom_field('Executive summary of the action')

    @property
    def cf_cross_cutting_issues(self):
        return self.get_custom_field('Cross-cutting issues')

    @property
    def cf_monitoring_evaluation(self):
        return self.get_custom_field('Monitoring & evaluation')

    @property
    def cf_learning_from_the_action(self):
        return self.get_custom_field('Learning from the action')

    @property
    def cf_list_of_materials_produced(self):
        return self.get_custom_field('List of materials produced')

    @property
    def cf_relationship_with_state_authorities(self):
        return self.get_custom_field('Relationship with State authorities')

    @property
    def cf_list_of_contracts(self):
        return self.get_custom_field('List of contracts')

    @property
    def cf_relationship_with_other_organisations(self):
        return self.get_custom_field('Relationship with other organisations')

    @property
    def cf_results_and_activities(self):
        return self.get_custom_field('Results and activities')

    def get_custom_field(self, name):
        if self._custom_fields is None:
            self._custom_fields = {f.name: f.value for f in self.custom_fields.all()}
        return self._custom_fields.get(name, '')


def build_view_object(project):
    periods = IndicatorPeriod.objects\
        .select_related('indicator', 'indicator__result', 'indicator__result__project')\
        .filter(indicator__result__project=project)

    if not periods.count():
        return EUTFProjectProxy(project)

    return utils.make_project_proxies(periods, EUTFProjectProxy)[0]


@login_required
@with_download_indicator
def render_report(request, project_id):
    project = get_object_or_404(Project, pk=project_id)

    project_view = build_view_object(project)

    if request.GET.get('show-html', ''):
        html = render_to_string('reports/eutf-narrative.html', context={'project': project_view})
        return HttpResponse(html)

    doc = Document(os.path.join(os.path.dirname(__file__), 'eutf-narrative.tpl.docx'))
    doc.sections[0].page_width = Mm(210)
    doc.sections[0].page_height = Mm(297)

    doc.add_heading("Narrative report", 0)

    # 1. Description
    table1 = doc.add_table(rows=11, cols=2)
    table1.style = 'Table with Paddings'
    table1.cell(0, 0).merge(table1.cell(0, 1))
    table1_title = table1.rows[0].cells[0].paragraphs[-1]
    table1_title.text = '1. Description'
    table1_title.style = 'Heading 1'

    table1.rows[1].cells[0].paragraphs[-1].add_run('1.1. Contact Person').bold = True
    table1.rows[1].cells[1].text = project_view.contact_person

    table1.rows[2].cells[0].paragraphs[-1].add_run('1.2. Partners').bold = True
    partners_count = project_view.all_partners().count()
    if partners_count:
        table1.rows[2].cells[1]._element.clear_content()
        partners_table = table1.rows[2].cells[1].add_table(rows=partners_count, cols=1)
        partners_table.style = 'Table Compact'
        for num, partner in enumerate(project_view.all_partners()):
            partners_table.rows[num].cells[0].text = partner.name

    table1.rows[3].cells[0].paragraphs[-1].add_run('1.3. Title of the actions').bold = True
    table1.rows[3].cells[1].text = project_view.title

    table1.rows[4].cells[0].paragraphs[-1].add_run('1.4. Contact number').bold = True
    table1.rows[4].cells[1].text = project_view.subtitle

    table1.rows[5].cells[0].paragraphs[-1].add_run('1.5. State date and end date of the action').bold = True
    table1.rows[5].cells[1].text = '{} - {}'.format(project_view.actual_start_year, project_view.planned_end_year)

    table1.rows[6].cells[0].paragraphs[-1].add_run('1.6. Target country(ies) or region(s)').bold = True
    recipient_countries_count = project_view.recipient_countries.count()
    if recipient_countries_count:
        table1.rows[6].cells[1]._element.clear_content()
        recipients_table = table1.rows[6].cells[1].add_table(rows=recipient_countries_count, cols=1)
        recipients_table.style = 'Table Compact'
        for num, country in enumerate(project_view.recipient_countries.all()):
            recipients_table.rows[num].cells[0].text = country.iati_country().name

    table1.rows[7].cells[0].paragraphs[-1].add_run('1.7. Final beneficiaries &/or target group').bold = True
    table1.rows[7].cells[1].text = project_view.target_group

    table1.rows[8].cells[0].paragraphs[-1].add_run('1.8. Co-funding').bold = True
    cofunding_partners_count = project_view.cofunding_partners.count()
    if cofunding_partners_count:
        table1.rows[8].cells[1]._element.clear_content()
        cofunding_partners_table = table1.rows[8].cells[1].add_table(rows=cofunding_partners_count, cols=2)
        cofunding_partners_table.style = 'Table Compact'
        for num, partner in enumerate(project_view.cofunding_partners):
            cofunding_partners_table.rows[num].cells[0].text = '{:,}'.format(partner.funding_amount)
            cofunding_partners_table.rows[num].cells[1].text = partner.organisation.name

    table1.rows[9].cells[0].paragraphs[-1].add_run('1.9. EUTF contribution').bold = True
    table1.rows[9].cells[1].text = project_view.eutf_funding_amount

    doc.add_paragraph('')

    # 2. Assessment of action activities
    table2 = doc.add_table(rows=8, cols=1)
    table2.style = 'Table with Paddings'
    table2_title = table2.rows[0].cells[0].paragraphs[-1]
    table2_title.text = '2. Assessment of action activities'
    table2_title.style = 'Heading 1'

    table2.rows[1].cells[0].paragraphs[-1].add_run('2.1. Executive summary of the action').bold = True
    markdown_to_docx(table2.rows[1].cells[0].paragraphs[-1], project_view.cf_executive_summary_of_the_action)

    table2.rows[2].cells[0].paragraphs[-1].add_run('2.2. Results and activities').bold = True
    markdown_to_docx(table2.rows[2].cells[0].paragraphs[-1], project_view.cf_results_and_activities)

    table2.rows[3].cells[0].paragraphs[-1].add_run('2.3. Sustainability').bold = True
    markdown_to_docx(table2.rows[3].cells[0].paragraphs[-1], project_view.cf_sustainability)

    table2.rows[4].cells[0].paragraphs[-1].add_run('2.4. Monitoring & evaluation').bold = True
    markdown_to_docx(table2.rows[4].cells[0].paragraphs[-1], project_view.cf_monitoring_evaluation)

    table2.rows[5].cells[0].paragraphs[-1].add_run('2.5. Learning from the action').bold = True
    markdown_to_docx(table2.rows[5].cells[0].paragraphs[-1], project_view.cf_learning_from_the_action)

    table2.rows[6].cells[0].paragraphs[-1].add_run('2.6. List of materials produced').bold = True
    markdown_to_docx(table2.rows[6].cells[0].paragraphs[-1], project_view.cf_list_of_materials_produced)

    table2.rows[7].cells[0].paragraphs[-1].add_run('2.7. List of contracts').bold = True
    markdown_to_docx(table2.rows[7].cells[0].paragraphs[-1], project_view.cf_list_of_contracts)

    doc.add_paragraph('')

    # 3. Beneficiaries/affiliated entities ...
    table3 = doc.add_table(rows=6, cols=1)
    table3.style = 'Table with Paddings'
    table3_title = table3.rows[0].cells[0].paragraphs[-1]
    table3_title.text = '3. Beneficiaries/affiliated entities and other Cooperation'
    table3_title.style = 'Heading 1'

    table3.rows[1].cells[0].paragraphs[-1].add_run('3.1. Relationship with beneficiaries').bold = True
    markdown_to_docx(table3.rows[1].cells[0].paragraphs[-1], project_view.cf_relationship_with_beneficiaries)

    table3.rows[2].cells[0].paragraphs[-1].add_run('3.2. Relationship with State authorities').bold = True
    markdown_to_docx(table3.rows[2].cells[0].paragraphs[-1], project_view.cf_relationship_with_state_authorities)

    table3.rows[3].cells[0].paragraphs[-1].add_run('3.3. Relationship with other organisations').bold = True
    markdown_to_docx(table3.rows[3].cells[0].paragraphs[-1], project_view.cf_relationship_with_other_organisations)

    table3.rows[4].cells[0].paragraphs[-1].add_run('3.4. Synergies with other actions').bold = True
    markdown_to_docx(table3.rows[4].cells[0].paragraphs[-1], project_view.cf_synergies_with_other_actions)

    table3.rows[5].cells[0].paragraphs[-1].add_run('3.5. Cooperation with contracting authority').bold = True
    markdown_to_docx(table3.rows[5].cells[0].paragraphs[-1], project_view.cf_cooperation_with_contracting_authority)

    doc.add_paragraph('')

    # 4. Visibility
    table4 = doc.add_table(rows=2, cols=1)
    table4.style = 'Table with Paddings'
    table4_title = table4.rows[0].cells[0].paragraphs[-1]
    table4_title.text = '4. Visibility'
    table4_title.style = 'Heading 1'

    table4.rows[1].cells[0].paragraphs[-1].add_run('4.1. EU visibility').bold = True
    markdown_to_docx(table4.rows[1].cells[0].paragraphs[-1], project_view.cf_eu_visibility)

    doc.add_paragraph('')

    # 5. Additional comments
    table5 = doc.add_table(rows=5, cols=2)
    table5.style = 'Table with Paddings'
    table5.cell(0, 0).merge(table5.cell(0, 1))
    table5_title = table5.rows[0].cells[0].paragraphs[-1]
    table5_title.text = '5. Additional comments'
    table5_title.style = 'Heading 1'

    table5.rows[1].cells[0].paragraphs[-1].add_run('5.1. Additional comments').bold = True
    markdown_to_docx(table5.rows[1].cells[1].paragraphs[-1], project_view.cf_additional_comments)

    table5.cell(2, 0).merge(table5.cell(2, 1))
    table5.rows[2].cells[0].text = (
        'Name of the contact person for the action: .................................................')

    table5.cell(3, 0).merge(table5.cell(3, 1))
    table5.rows[3].cells[0].text = (
        'Signature: ................................................. '
        'Location: .................................................')

    table5.cell(4, 0).merge(table5.cell(4, 1))
    table5.rows[4].cells[0].text = (
        'Date report due:  ................................................. '
        'Date report sent: .................................................')

    doc.add_paragraph('')

    change_orientation(doc)

    doc.add_heading('Appendix 1 - Logframe', 2)
    log_table = doc.add_table(rows=2, cols=9)
    log_table.style = 'Table with Grid'
    log_table.rows[0].cells[0].paragraphs[-1].add_run('Type').bold = True
    log_table.rows[0].cells[0].width = Mm(20)
    log_table.rows[0].cells[1].paragraphs[-1].add_run('Results').bold = True
    log_table.rows[0].cells[1].width = Mm(50)
    log_table.rows[0].cells[2].paragraphs[-1].add_run('Indicators').bold = True
    log_table.rows[0].cells[4].width = Mm(50)
    log_table.rows[0].cells[3].paragraphs[-1].add_run('Baseline').bold = True
    log_table.rows[0].cells[3].width = Mm(40)
    log_table.rows[0].cells[5].paragraphs[-1].add_run('Target').bold = True
    log_table.rows[0].cells[5].width = Mm(40)
    log_table.rows[0].cells[7].paragraphs[-1].add_run('Actual').bold = True
    log_table.rows[0].cells[7].width = Mm(20)
    log_table.rows[0].cells[8].paragraphs[-1].add_run('Comments').bold = True
    log_table.rows[0].cells[8].width = Mm(50)
    log_table.cell(0, 3).merge(log_table.cell(0, 4))
    log_table.cell(0, 5).merge(log_table.cell(0, 6))

    log_table.cell(1, 0).merge(log_table.cell(1, 2))
    log_table.rows[1].cells[0].width = Mm(120)
    log_table.rows[1].cells[3].paragraphs[-1].add_run('Year').bold = True
    log_table.rows[1].cells[3].width = Mm(20)
    log_table.rows[1].cells[4].paragraphs[-1].add_run('Value').bold = True
    log_table.rows[1].cells[4].width = Mm(20)
    log_table.rows[1].cells[5].paragraphs[-1].add_run('Year').bold = True
    log_table.rows[1].cells[5].width = Mm(20)
    log_table.rows[1].cells[6].paragraphs[-1].add_run('Value').bold = True
    log_table.rows[1].cells[6].width = Mm(20)
    log_table.rows[1].cells[7].paragraphs[-1].add_run('Value').bold = True
    log_table.rows[1].cells[7].width = Mm(20)
    log_table.rows[1].cells[8].width = Mm(50)

    set_repeat_table_header(log_table.rows[0])
    set_repeat_table_header(log_table.rows[1])

    for result in project_view.results:
        for indicator in result.indicators:
            for period in indicator.periods:
                row = log_table.add_row()
                row.cells[0].text = result.iati_type_name
                row.cells[0].width = Mm(20)
                row.cells[1].text = result.title
                row.cells[1].width = Mm(50)
                row.cells[2].text = indicator.title
                row.cells[2].width = Mm(50)
                row.cells[3].text = indicator.baseline_year or ''
                row.cells[3].width = Mm(20)
                row.cells[4].text = indicator.baseline_value or ''
                row.cells[4].width = Mm(20)
                if period.period_start:
                    row.cells[5].text = period.period_start.strftime('%Y')
                row.cells[5].width = Mm(20)
                if period.target_value:
                    row.cells[6].text = '{:,}'.format(period.target_value)
                row.cells[6].width = Mm(20)
                if period.actual_value:
                    row.cells[7].text = '{:,}'.format(period.actual_value)
                row.cells[7].width = Mm(20)
                row.cells[8].text = period.actual_comment
                row.cells[8].width = Mm(50)

    filename = '{}-{}-eutf-narrative.docx'.format(datetime.today().strftime('%Y%b%d'), project_view.id)

    return utils.make_docx_response(doc, filename)
